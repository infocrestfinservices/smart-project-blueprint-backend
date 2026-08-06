"""
word_builder.py

Turns the structured model JSON into a professional, purpose-specific .docx
(python-docx). Structure (sections, terminology) is driven by purpose_config;
narrative prose comes from the AI.

The financial figures shown here are read from `model["financial_summary"]`, which
is extracted from the SAME recalculated Excel workbook the user downloads — so the
Word report's numbers are the Excel's numbers.

DESIGN
------
Layout follows the client's reference report: a cover with a full-height indigo
band down the left edge, large NUMBERED section headings set in indigo (plain
coloured type, not filled bars), tables with an indigo header row and alternating
lavender rows, indigo charts, and a page number in the footer. Only the look is
borrowed — every word and figure here is this client's own.
"""

import io
import logging
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from purpose_config import get_config

logger = logging.getLogger("word_builder")

# ── palette (from the reference report) ─────────────────────────────────────────
INDIGO = RGBColor(0x5B, 0x5B, 0xF5)
INDIGO_HEX = "5B5BF5"
INDIGO_DEEP = RGBColor(0x2B, 0x2B, 0xD4)
INDIGO_DEEP_HEX = "2B2BD4"
LAVENDER = "EFEFFB"          # alternating table rows
LAVENDER_SOFT = "F6F6FD"
INK = RGBColor(0x1C, 0x1C, 0x22)
GREY = RGBColor(0x6B, 0x70, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── formatting helpers ──────────────────────────────────────────────────────────
def _fmt_date(d):
    try:
        dt = datetime.fromisoformat(str(d).replace("Z", "")) if d else datetime.utcnow()
    except (ValueError, TypeError):
        dt = datetime.utcnow()
    return dt.strftime("%d %B %Y")


def _inr(v):
    if not isinstance(v, (int, float)):
        return "—"
    a = abs(v)
    if a >= 1e7:
        return f"₹{v / 1e7:,.2f} Cr"
    if a >= 1e5:
        return f"₹{v / 1e5:,.2f} L"
    return f"₹{v:,.0f}"


def _ratio(v):
    return f"{v:.2f}×" if isinstance(v, (int, float)) else "—"


def _pct(v):
    return f"{v * 100:.1f}%" if isinstance(v, (int, float)) else "—"


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.25
    # Narrow margins so the cover's colour band reads as a full-height edge band
    # rather than a floating rectangle.
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.8)


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _borders(table, *, hex_color="FFFFFF", size=8, inside=True):
    """The reference tables have no visible grid — rows are separated by fill, with a
    thin white gutter. Setting white borders reproduces that."""
    tblPr = table._tbl.tblPr
    el = OxmlElement("w:tblBorders")
    edges = ("top", "left", "bottom", "right") + (("insideH", "insideV") if inside else ())
    for edge in edges:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), hex_color)
        el.append(e)
    tblPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        el.append(e)
    tblPr.append(el)


def _bookmark(paragraph, name, bid):
    """Mark a paragraph so an internal hyperlink can jump to it."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    p = paragraph._p
    # bookmarkStart has to sit AFTER pPr — pPr must be the first child or Word rejects
    # the paragraph outright.
    p.insert(1 if p.find(qn("w:pPr")) is not None else 0, start)
    p.append(end)


def _anchor_name(doc):
    """Next bookmark name for this document. Word's own contents uses _Toc… names."""
    n = getattr(doc, "_bm_seq", 0) + 1
    doc._bm_seq = n
    return f"_Toc{n:05d}", n


def _read_anchor(paragraph):
    bs = paragraph._p.find(qn("w:bookmarkStart"))
    return bs.get(qn("w:name")) if bs is not None else None


def _internal_link(paragraph, anchor, text):
    """A run that jumps to a bookmark in the same document."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    run = paragraph.add_run(text)
    link.append(run._r)               # move the run inside the hyperlink
    paragraph._p.append(link)
    return run


def _external_link(paragraph, url, text):
    """A run that opens a URL. Needs a relationship on the containing part."""
    rid = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = paragraph.add_run(text)
    link.append(run._r)
    paragraph._p.append(link)
    return run


def _cell_text(cell, text, *, bold=False, color=None, size=10, align=None, white=False,
               space=3, anchor=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(space)
    p.paragraph_format.space_after = Pt(space)
    if align:
        p.alignment = align
    run = _internal_link(p, anchor, str(text)) if anchor else p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = WHITE
    elif color is not None:
        run.font.color.rgb = color


def _page_footer(doc, title="", logo=None):
    """Footer with "Page X of Y" bottom-right, and a header carrying the report title and,
    when the client has uploaded one, their logo top-right — on every page except the
    cover (the first-page header/footer are left blank)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True     # keep the cover clean
    # Explicitly give the cover (first page) its own EMPTY header/footer so no page number
    # or title bleaks onto it — some renderers ignore the flag unless the parts exist.
    for part in (section.first_page_header, section.first_page_footer):
        part.is_linked_to_previous = False
        if not part.paragraphs:
            part.add_paragraph()
        part.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = fp.add_run("Page "); r0.font.size = Pt(9); r0.font.color.rgb = GREY
    _add_field(fp, "PAGE")
    r1 = fp.add_run(" of "); r1.font.size = Pt(9); r1.font.color.rgb = GREY
    _add_field(fp, "NUMPAGES")
    for r in fp.runs:
        r.font.size = Pt(9); r.font.color.rgb = GREY

    if title or logo:
        header = section.header
        hp = header.paragraphs[0]
        if logo:
            # The client's mark, top-right on every page after the cover — a two-cell
            # table so the title can sit left while the logo holds the right edge, which
            # a single paragraph cannot do.
            hp.text = ""
            t = header.add_table(rows=1, cols=2, width=Cm(17.0))
            _no_borders(t)
            left, right = t.rows[0].cells
            left.width, right.width = Cm(13.0), Cm(4.0)
            lp = left.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if title:
                lr = lp.add_run(title)
                lr.font.size = Pt(8); lr.italic = True; lr.font.color.rgb = GREY
            rp = right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                rp.add_run().add_picture(io.BytesIO(logo), height=Cm(0.9))
            except Exception:
                logger.warning("header: could not place the logo", exc_info=True)
            rule_holder = header.add_paragraph()
            _bottom_rule(rule_holder, color="D9D9E6", size=4)
        else:
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hr = hp.add_run(title)
            hr.font.size = Pt(8); hr.italic = True; hr.font.color.rgb = GREY
            _bottom_rule(hp, color="D9D9E6", size=4)


# ── fields & rules (used by the TOC / captions / heading underline) ─────────────
def _add_field(paragraph, instr, hint=""):
    """Insert a Word field (TOC, SEQ, PAGE …). It stays a live field, so Word fills in
    the real value/page numbers when the document is opened (build_word sets the doc to
    update fields on open). `hint` is the placeholder shown until that happens."""
    r = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    r.append(begin); r.append(it); r.append(sep)
    if hint:
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = hint; r.append(t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); r.append(end)


def _shade_paragraph(paragraph, hex_fill):
    """A tinted band behind a single paragraph — the cover's proposal ribbon."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _bottom_rule(paragraph, color=INDIGO_HEX, size=6):
    """A thin coloured line under a paragraph — used to underline main headings."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)


# ── section heading — three levels, always bold, sized 20 / 16 / 12 ─────────────
class _Numbering:
    def __init__(self):
        self.n = 0
        self.entries = []      # kept for compatibility; the Contents is now a live field

    def next(self):
        self.n += 1
        return self.n


# level -> (point size, colour, space-before). Consistent 4-pt step, per the brief.
_HEADING_LEVELS = {
    1: (20, INDIGO, 22),
    2: (16, INDIGO_DEEP, 14),
    3: (12, INK, 10),
}


# Everything the client inserts is set to ONE width, whatever the file's own size — a
# report where each photograph is a different width reads as a scrapbook. Height follows
# the aspect ratio, so nothing is squashed.
INSERT_WIDTH_IN = 5.6


def _set_inserts(doc, inserts):
    """Attach the client's own images, keyed by the section each belongs at the end of.

    `inserts` is [{section, image: bytes, caption}]. Held on the document because a
    section's images can only be placed once the section is finished, and the composer
    walks it linearly.
    """
    doc._inserts = {}
    for it in inserts or []:
        section = str((it or {}).get("section") or "").strip()
        if section and it.get("image"):
            doc._inserts.setdefault(section, []).append(it)
    doc._cur_section = None
    doc._section_titles = []


def _flush_inserts(doc):
    """Place the images belonging to the section that just ended."""
    pending = getattr(doc, "_inserts", None)
    current = getattr(doc, "_cur_section", None)
    if not pending or not current:
        return
    for it in pending.pop(current, []):
        try:
            _add_image(doc, it["image"], INSERT_WIDTH_IN)
        except Exception:
            logger.warning("insert: could not place an image in %r", current, exc_info=True)
            continue
        if str(it.get("caption") or "").strip():
            # Captioned through the normal helper, so a client's photograph is numbered
            # and listed in the List of Figures alongside the generated charts.
            _figure_caption(doc, str(it["caption"]).strip())


def _heading(doc, text, numbering=None, *, sub=None, new_page=False, level=1):
    # A main heading ends the previous section, which is the moment its images go in.
    if level == 1:
        _flush_inserts(doc)
        doc._cur_section = text
        # The BARE title, recorded as the sections are raised. Reading them back off the
        # rendered headings instead would give "3. Business Model", and an insert saved
        # against that would match no section — the numbering is added at render time.
        doc._section_titles.append(text)
    p = doc.add_paragraph()
    # Use the built-in Heading style so the Table of Contents field can pick the heading
    # up at the right outline level; the run font below overrides its look.
    try:
        p.style = doc.styles[f"Heading {level}"]
    except (KeyError, ValueError):
        pass
    # A main section ALWAYS opens a page of its own, and no heading of any level may be
    # left stranded at the foot of a page with its text overleaf: keep_with_next binds a
    # heading to what follows it, so the renderer carries it forward instead.
    if new_page or level == 1:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    size, colour, before = _HEADING_LEVELS.get(level, _HEADING_LEVELS[1])
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _bookmark(p, *_anchor_name(doc))          # so the contents can link to it
    if numbering is not None and level == 1:
        n = numbering.next()
        numbering.entries.append((n, text))
        label = f"{n}. {text}"
    else:
        label = text
    run = p.add_run(label)
    run.bold = True                       # headings are ALWAYS bold
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.font.name = "Calibri"
    if level == 1:
        _bottom_rule(p)                   # a thin rule under every main heading
    if sub:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        sp.paragraph_format.keep_with_next = True   # the sub-line follows its heading
        sr = sp.add_run(sub)
        sr.font.size = Pt(9)
        sr.italic = True
        sr.font.color.rgb = GREY
    return p


def _table_caption(doc, text):
    """A numbered "Table N:  …" caption above each table. The SEQ field auto-numbers it
    and lets the List of Tables collect it (TOC \\c "Table")."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True     # never strand a caption from its table
    _bookmark(p, *_anchor_name(doc))
    r0 = p.add_run("Table ")
    r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = INK
    _add_field(p, "SEQ Table \\* ARABIC")
    r1 = p.add_run(f":  {text}")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = INK


def _figure_caption(doc, text):
    """A numbered "Figure N:  …" caption for a chart, collected by the List of Figures."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    _bookmark(p, *_anchor_name(doc))
    r0 = p.add_run("Figure ")
    r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = GREY
    _add_field(p, "SEQ Figure \\* ARABIC")
    r1 = p.add_run(f":  {text}")
    r1.font.size = Pt(10); r1.font.color.rgb = GREY


def _cut_hold(paragraph):
    """End a keep-with-next run at the paragraph before a heading.

    A heading carries keep_with_next, and so do the two paragraphs under it, so that a
    sub-section never begins with a single orphan line at the foot of a page. But left
    unbroken those links chain — heading, paragraph, NEXT heading, its paragraph — and
    once the chain is taller than the space left, the renderer abandons it and drops the
    whole run to the next page, leaving the first heading stranded exactly as before.
    Cutting the link on the paragraph that precedes a heading keeps each run short enough
    to actually be honoured.
    """
    if paragraph is not None:
        paragraph.paragraph_format.keep_with_next = False


def _render_markdownish(doc, text):
    hold, held = 0, None
    for raw in str(text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            head = line.lstrip("#").strip().lstrip("*").strip("* ").strip()
            if head:
                hashes = len(line) - len(line.lstrip("#"))
                _cut_hold(held)
                _heading(doc, head, level=3 if hashes >= 3 else 2)   # sub-heading -> Contents
                hold, held = 2, None
            continue
        if line.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        held = None
        if hold:
            p.paragraph_format.keep_with_next = True
            held = p
            hold -= 1
    _cut_hold(held)


def _add_image(doc, png_bytes, width_in):
    if not png_bytes:
        return
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fixed_columns(table, cms):
    """Pin column widths. A renderer only honours these under a FIXED layout with an
    explicit grid — otherwise it autofits and redistributes them."""
    table.autofit = False
    for cell, cm in zip(table.rows[0].cells, cms):
        cell.width = Cm(cm)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, cm in zip(grid.findall(qn("w:gridCol")), cms):
            gc.set(qn("w:w"), str(int(cm * 567)))


def _caption_list(doc):
    """(kind, title, anchor) for every numbered caption, in order — what the two lists show.

    The SEQ field's number is not in the paragraph text until a renderer fills it, so the
    captions are numbered here in document order, which is what SEQ produces anyway.
    """
    out = []
    for p in doc.paragraphs:
        # "Table <SEQ>:  Title" — the SEQ field renders as nothing here, hence \d* .
        m = re.match(r"^(Table|Figure)\s*\d*\s*:\s*(.+)$", p.text.strip())
        if m:
            out.append((m.group(1), m.group(2).strip(), _read_anchor(p)))
    return out


def _list_block(doc, heading, col_label, rows, bold_top=True):
    """One of the three front-matter lists, drawn as a two-column table.

    `rows` is [(level, text, page, anchor)]; a page of None prints blank, which is what the
    measuring pass needs so it paginates the same as the final document, and an anchor
    makes the entry a working link to that place in the document.
    """
    if not rows:
        return
    h = doc.add_paragraph()
    h.paragraph_format.page_break_before = True
    h.paragraph_format.space_after = Pt(8)
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(20)
    hr.font.color.rgb = INDIGO
    hr.font.name = "Calibri"
    _bottom_rule(h)

    table = doc.add_table(rows=1, cols=2)
    _borders(table)
    # Without a fixed grid Word AUTOFITS to the content, and the contents entries are
    # short ("1. Executive Summary"), so the whole table shrank and the right-aligned page
    # number landed in the middle of the sheet instead of at the margin. The two caption
    # lists looked fine only because their labels are long enough to fill the width.
    COLS = (15.1, 2.5)          # A4 21.0 less the 1.6/1.8 margins = 17.6 cm of text
    _fixed_columns(table, COLS)
    head_cells = table.rows[0].cells
    _cell_text(head_cells[0], col_label, bold=True, white=True, size=9.5)
    # The page column in the front matter is right-aligned — asked for explicitly, and
    # deliberately NOT routed through FIGURE_ALIGN, which governs the data tables.
    _cell_text(head_cells[1], "Page", bold=True, white=True, size=9.5,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    for c in head_cells:
        _shade(c, INDIGO_HEX)
    for band, (level, text, page, anchor) in enumerate(rows):
        top = bold_top and level == 1
        cells = table.add_row().cells
        for c, cm in zip(cells, COLS):
            c.width = Cm(cm)
        # Both cells link, so clicking anywhere on the line jumps to the section. The
        # entries keep the contents' own colour rather than the blue-underline web look.
        _cell_text(cells[0], ("    " * (level - 1)) + text,
                   size=10 if top else 9.5, bold=top,
                   color=INK if top else GREY, anchor=anchor)
        _cell_text(cells[1], "" if page is None else str(page),
                   size=10 if top else 9.5, bold=top,
                   color=INK if top else GREY,
                   align=WD_ALIGN_PARAGRAPH.RIGHT, anchor=anchor)
        for c in cells:
            _shade(c, LAVENDER if band % 2 == 0 else LAVENDER_SOFT)
    doc.add_paragraph()


def _toc_pages(doc, entries=None, figures=None, tables=None):
    """The three front-matter lists, written out with real page numbers.

    A live `TOC` field depends on the reader's Word repaginating on open, and it renders
    as nothing at all through headless LibreOffice (which is how the PDF download is
    produced). In practice every entry came out as page 1. So the pages are measured from
    a real render (see _measure_pages) and typed in — works everywhere, no F9.

    Each argument is [(level, text, page)]. A page of None prints blank, so the measuring
    pass paginates identically to the final document.
    """
    _list_block(doc, "Table of Contents", "Section", entries)
    _list_block(doc, "List of Figures", "Figure", figures, bold_top=False)
    _list_block(doc, "List of Tables", "Table", tables, bold_top=False)


def _heading_list(doc):
    """(level, text, anchor) for every real heading, in document order."""
    out = []
    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        if not style.startswith("Heading"):
            continue
        try:
            level = int(style.split()[-1])
        except (ValueError, IndexError):
            continue
        text = p.text.strip()
        if level <= 3 and text:
            out.append((level, text, _read_anchor(p)))
    return out


_FRONT_MATTER = ("Table of Contents", "List of Figures", "List of Tables")


def _measure_pages(docx_bytes, headings, captions=()):
    """Render the document and read back the page each heading and caption landed on.

    `headings` is [(level, text, anchor)] and `captions` is [(kind, title, anchor)], both in
    DOCUMENT order — figures and tables interleave, so they are located in one pass;
    searching the figures and then the tables would make the monotonic scan below skip
    every table that sits before the last figure.

    Returns ([page per heading], [page per caption]). Without LibreOffice, or if
    anything about the render is unexpected, returns empty results and the lists are
    written without numbers rather than with wrong ones.
    """
    blank = ([None] * len(headings), [None] * len(captions))
    try:
        from services.recalc_service import libreoffice_available
        if not libreoffice_available():
            return blank
        import fitz
        pdf = _docx_to_pdf(docx_bytes)
        if not pdf:
            return blank
        with fitz.open(stream=pdf, filetype="pdf") as d:
            texts = [d[i].get_text() for i in range(d.page_count)]

        # The front matter lists every heading and every caption, so a naive search finds
        # them all there. The discriminator is DENSITY, but the threshold has to clear the
        # densest real section: a page of sub-headings had seven, while a front-matter page
        # carries dozens. Anything at or above this is front matter, whether it is drawn as
        # tabbed lines or as a table (a table has no dotted leader to look for). The three
        # literal titles are checked as well, so a short list is never mistaken for body.
        probes = [t[:60] for _, t, _ in headings] + [t[:60] for _, t, _ in captions]
        density = [sum(1 for pr in probes if pr and pr in body) for body in texts]
        body_max = max((n for n in density if n < 12), default=0)
        cutoff = max(12, body_max + 4)
        front = {i for i, n in enumerate(density)
                 if n >= cutoff or any(t in texts[i] for t in _FRONT_MATTER)}

        def _locate(items):
            # Items appear in document order, so each search starts where the last one
            # landed — that stops a title mentioned in earlier prose from stealing the page.
            found, start = [], 0
            for text in items:
                probe = text[:60]
                page = None
                for i in range(start, len(texts)):
                    if i in front:
                        continue
                    if probe and probe in texts[i]:
                        page, start = i + 1, i
                        break
                found.append(page)
            return found

        return (_locate([t for _, t, _ in headings]),
                _locate([t for _, t, _ in captions]))
    except Exception:
        logger.warning("contents: could not measure page numbers", exc_info=True)
        return blank


def _docx_to_pdf(docx_bytes):
    """LibreOffice headless docx -> pdf. None on any failure.

    Uses recalc_service's binary discovery — `soffice` is usually not on PATH on Windows,
    which is why calling it by bare name silently produced nothing.
    """
    import subprocess, tempfile, glob as _glob, os as _os
    from services.recalc_service import _soffice_bin
    exe = _soffice_bin()
    if not exe:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        src = _os.path.join(tmp, "doc.docx")
        with open(src, "wb") as fh:
            fh.write(docx_bytes)
        try:
            subprocess.run([exe, "--headless", "--convert-to", "pdf", "--outdir", tmp, src],
                           check=True, capture_output=True, timeout=240)
        except Exception:
            logger.warning("contents: docx -> pdf failed", exc_info=True)
            return None
        hits = _glob.glob(_os.path.join(tmp, "*.pdf"))
        if hits:
            with open(hits[0], "rb") as fh:
                return fh.read()
    return None


# ── cover ───────────────────────────────────────────────────────────────────────
def _cover(doc, config, project):
    """Full-height indigo band down the left edge; title block on the right."""
    table = doc.add_table(rows=1, cols=2)
    _no_borders(table)
    table.autofit = False
    band, body = table.rows[0].cells
    BAND_CM, BODY_CM = 3.4, 13.6
    band.width = Cm(BAND_CM)
    body.width = Cm(BODY_CM)
    # Word/LibreOffice only honours these widths under a FIXED layout with an explicit
    # grid — without it the renderer re-distributes the columns (the band came out
    # ~9 cm wide and the artwork was clipped by the squeezed title column).
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, cm in zip(grid.findall(qn("w:gridCol")), (BAND_CM, BODY_CM)):
            gc.set(qn("w:w"), str(int(cm * 567)))
    _shade(band, INDIGO_HEX)
    band.text = ""
    # EXACT height: without the rule the cell grows past the page and spills a blank
    # sheet before the content.
    row = table.rows[0]
    row.height = Cm(25.6)
    trPr = row._tr.get_or_add_trPr()
    rh = OxmlElement("w:trHeight")
    rh.set(qn("w:val"), str(int(25.6 * 567)))
    rh.set(qn("w:hRule"), "exact")
    trPr.append(rh)

    # Everything on the cover reads left-to-right off the indigo band, so LEFT is the
    # default. The promoter line and logo used to be right-aligned, which left the cover
    # ragged on the wrong edge against a left-hand band.
    def para(text, *, size=11, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT,
             italic=False, space_before=0, space_after=6):
        p = body.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        return p

    # ── the band: a monogram plate, the business name, and a confidential mark ──────
    # An empty colour bar is just decoration; carrying the mark turns it into the
    # report's identity, which is what makes the page memorable at a glance.
    title_text = str(project.get("title") or "Project Report")
    initials = "".join(w[0] for w in re.findall(r"[A-Za-z]+", title_text)[:2]).upper() or "R"
    logo = (project.get("branding") or {}).get("logo")

    def band_para(text="", *, size=9, bold=False, space_before=0, space_after=4,
                  color=WHITE, spacing=True):
        p = band.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(" ".join(text) if spacing else text)
            r.bold = bold
            r.font.size = Pt(size)
            r.font.color.rgb = color
        return p

    band.paragraphs[0].text = ""
    for _ in range(3):
        band_para()
    if logo:
        # Their mark takes the plate when they have one; otherwise the initials do.
        try:
            lp = band_para(space_before=6, space_after=6)
            lp.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))
        except Exception:
            logger.warning("cover: could not place the uploaded logo", exc_info=True)
            logo = None
    if not logo:
        _bottom_rule(band_para(space_after=10), color="FFFFFF", size=8)
        band_para(initials, size=26, bold=True, space_after=10, spacing=False)
        _bottom_rule(band_para(space_after=10), color="FFFFFF", size=8)
    # The name goes one WORD per line. Letter-spacing a long title across a 3.4 cm band
    # split it mid-word ("MANGO PICKLE MA / NUFACTURING UNI / T"), which looks like a
    # rendering fault rather than a design.
    for word in title_text.upper().split()[:4]:
        band_para(word if len(word) > 11 else " ".join(word),
                  size=7.5, bold=True, space_after=1, spacing=False)
    for _ in range(14):
        band_para(space_after=0)
    _bottom_rule(band_para(space_after=6), color="FFFFFF", size=4)
    band_para("CONFIDENTIAL", size=7, bold=True, space_after=2)
    band_para(f"{config['label'].upper()} REPORT", size=7, bold=True, space_after=0)

    # ── the body ───────────────────────────────────────────────────────────────────
    body.paragraphs[0].text = ""
    promoter = project.get("promoter_name")
    if promoter:
        para("PROMOTER", size=8.5, bold=True, color=INDIGO, space_after=1)
        para(str(promoter), size=13, bold=True, space_after=0)
    for _ in range(3):
        body.add_paragraph()
    # A tinted ribbon naming the ask — the one line a banker looks for first. It uses the
    # PURPOSE ("bank_loan" -> "BANK LOAN PROPOSAL"), not the report label, which would
    # only repeat the "CMA DATA" eyebrow directly beneath it.
    ask = str(project.get("purpose") or "").replace("_", " ").strip()
    ribbon_text = (f"{ask} Proposal" if ask else f"{config['label']} Proposal").upper()
    ribbon = para(ribbon_text, size=11, bold=True, color=INDIGO_DEEP,
                  space_before=2, space_after=2)
    _shade_paragraph(ribbon, "ECECFB")
    body.add_paragraph()
    # eyebrow label (report type) — small, tracked, above the business name
    para(config["label"].upper(), size=11, bold=True, color=INDIGO,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    title_p = para(title_text, size=30, bold=True, color=INK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    _bottom_rule(title_p, size=14)      # a solid rule directly under the title, not floating
    body.add_paragraph()
    # a clean fact block, left-aligned, quiet grey — consistent vertical rhythm
    facts = []
    if project.get("industry"):
        facts.append(("Industry", project["industry"]))
    loc = project.get("location") or project.get("country")
    if loc:
        facts.append(("Location", loc))
    facts.append(("Prepared", _fmt_date(project.get("created_date") or project.get("created_at"))))
    bank = project.get("bank_name") or (project.get("purpose_answers") or {}).get("bank_name")
    if bank:
        facts.append(("Bank", str(bank)))
    for k, v in facts:
        fp = body.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.paragraph_format.space_after = Pt(3)
        rk = fp.add_run(f"{k}   ")
        rk.bold = True; rk.font.size = Pt(9.5); rk.font.color.rgb = INDIGO_DEEP
        rv = fp.add_run(str(v))
        rv.font.size = Pt(10.5); rv.font.color.rgb = GREY
    # Industry artwork fills the empty lower half of the cover. It is drawn in the
    # report's own palette and anchored to the bottom of the title column, so it reads
    # as part of the cover rather than a picture dropped onto it.
    art = None
    try:
        from services.cover_art import cover_art
        art = cover_art(project.get("industry"), project)
    except Exception:
        logger.warning("cover art unavailable", exc_info=True)
    if art:
        _bottom_rule(body.add_paragraph(), color="D9D9E6", size=4)
        holder = body.add_paragraph()
        holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        holder.paragraph_format.space_before = Pt(10)
        holder.paragraph_format.space_after = Pt(4)
        holder.add_run().add_picture(io.BytesIO(art), width=Inches(5.15))
        _bottom_rule(body.add_paragraph(), color="D9D9E6", size=4)
    else:
        for _ in range(3):
            body.add_paragraph()
    para("Prepared in accordance with professional Chartered Accountancy standards",
         size=9, italic=True, color=GREY)


# ── KPI tiles ──────────────────────────────────────────────────────────────────
def _kpi_cards(doc, cards):
    items = [
        ("Revenue · Year 5", _inr(cards.get("revenue_y5"))),
        ("EBITDA · Year 5", _inr(cards.get("ebitda_y5"))),
        ("PAT · Year 5", _inr(cards.get("pat_y5"))),
        ("Average DSCR", _ratio(cards.get("avg_dscr"))),
    ]
    items = [(l, v) for l, v in items if v != "—"]
    if not items:
        return
    table = doc.add_table(rows=2, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(table, hex_color="FFFFFF", size=14)     # a wider white gutter between tiles
    for i, (label, value) in enumerate(items):
        top, bot = table.cell(0, i), table.cell(1, i)
        _shade(top, INDIGO_HEX)
        _shade(bot, LAVENDER)
        _cell_text(top, value, bold=True, size=19, align=WD_ALIGN_PARAGRAPH.CENTER,
                   white=True, space=9)
        _cell_text(bot, label.upper(), size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=INDIGO_DEEP, space=5)
    doc.add_paragraph()


# ── tables ────────────────────────────────────────────────────────────────────
_NUMERIC = re.compile(r"^[\s(₹$€£]*-?[\d,]+(\.\d+)?\s*[%x×]?[\s)]*$", re.I)

# How a column of figures is set. The CA convention is right-aligned so digits line up,
# but the client reviewed it in place and asked for everything to read from the left
# edge — so figures and words start at the same margin. One constant, because this is
# exactly the sort of thing that gets reversed again.
FIGURE_ALIGN = WD_ALIGN_PARAGRAPH.LEFT


def _numeric_columns(headers, rows):
    """Which columns hold figures.

    Everything except the first column used to be right-aligned, so a sentence like
    "Supply to mandi, sub-distributors and grocery supply chains" came out ragged on its
    left edge — unreadable in a table of prose. Figures still belong on the right so their
    digits line up; words belong on the left.
    """
    out = []
    for i in range(len(headers)):
        vals = [str(r[i]).strip() for r in rows
                if i < len(r) and str(r[i]).strip() not in ("", "-", "—")]
        out.append(bool(vals) and sum(bool(_NUMERIC.match(v)) for v in vals) >= len(vals) * 0.6)
    return out


def _data_table(doc, headers, rows, *, first_align_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    _borders(table)
    numeric = _numeric_columns(headers, rows)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell_text(hdr[i], h, bold=True, white=True, size=9.5,
                   align=(FIGURE_ALIGN if numeric[i]
                          else None if (i == 0 and first_align_left)
                          else WD_ALIGN_PARAGRAPH.LEFT))
        _shade(hdr[i], INDIGO_HEX)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _cell_text(cells[i], val, size=9.5,
                       align=(FIGURE_ALIGN if numeric[i]
                              else None if (i == 0 and first_align_left)
                              else WD_ALIGN_PARAGRAPH.LEFT))
        fill = LAVENDER if ri % 2 == 0 else LAVENDER_SOFT
        for c in cells:
            _shade(c, fill)
    doc.add_paragraph()
    return table


def _charts_annexure(doc, summary, numbering, segments=None):
    """All charts together at the END of the report. Each sits on its own block with a
    caption; previously they were interleaved with the tables and split across pages,
    which is what made the layout look broken."""
    try:
        from services.report_charts import (revenue_profit_chart, dscr_chart, margin_chart,
                                            segment_donut, cost_structure_donut)
    except Exception:
        logger.warning("charts unavailable (matplotlib?); skipping annexure", exc_info=True)
        return
    charts = [
        (revenue_profit_chart(summary),
         "Revenue, EBITDA and Profit After Tax — five-year projection", 6.3),
        (dscr_chart(summary), "Debt Service Coverage Ratio by year", 4.6),
        (margin_chart(summary), "EBITDA and net profit margins by year", 4.6),
        (cost_structure_donut(summary), "Where each rupee of Year-1 revenue goes", 4.8),
        (segment_donut(segments), "Revenue mix across the target market segments", 4.8),
    ]
    charts = [(img, cap) for img, cap, _w in charts if img]
    if not charts:
        return
    _heading(doc, "Graphs & Charts", numbering,
             sub="Visual summary of the projections set out in the statements above",
             new_page=True)
    # Two per page at one common width — laid out as a single-column table so the
    # caption and its chart cannot drift apart or land at different sizes.
    CHART_W = 5.9
    for i, (img, caption) in enumerate(charts):
        if i and i % 2 == 0:
            doc.add_page_break()
        _figure_caption(doc, caption)            # "Figure N: …" -> List of Figures
        holder = doc.add_table(rows=1, cols=1)
        holder.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(holder)
        img_cell = holder.cell(0, 0)
        img_cell.text = ""
        run = img_cell.paragraphs[0].add_run()
        run.add_picture(io.BytesIO(img), width=Inches(CHART_W))
        img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    _charts_summary(doc, summary, segments)


def _charts_summary(doc, summary, segments):
    """A short read-out of what the charts show, drawn from the same figures."""
    cards = (summary or {}).get("cards") or {}
    series = (summary or {}).get("series") or {}
    ratios = (summary or {}).get("ratios") or {}
    lines = []

    rev = [v for v in series.get("Net Sales / Revenue", []) if isinstance(v, (int, float))]
    if len(rev) >= 2 and rev[0]:
        lines.append(f"Revenue rises from {_inr(rev[0])} in Year 1 to {_inr(rev[-1])} in "
                     f"Year 5, a {((rev[-1] / rev[0]) - 1) * 100:+.0f}% movement over the "
                     f"projection, with EBITDA and profit after tax widening alongside it.")
    dscr = [v for v in ratios.get("DSCR", []) if isinstance(v, (int, float))]
    if dscr:
        weak = [i + 1 for i, v in enumerate(dscr) if v < 1.20]
        avg = sum(dscr) / len(dscr)
        if weak:
            lines.append(f"Debt service coverage averages {avg:.2f} times. It sits below "
                         f"the 1.20 benchmark in Year {', '.join(map(str, weak))} while the "
                         f"unit is still building up, and clears it from Year "
                         f"{max(weak) + 1} onwards.")
        else:
            lines.append(f"Debt service coverage averages {avg:.2f} times and stays above "
                         f"the 1.20 benchmark throughout.")
    npm = [v for v in ratios.get("Net Profit Margin", []) if isinstance(v, (int, float))]
    if len(npm) >= 2:
        lines.append(f"Margins improve as fixed costs are spread over a larger turnover — "
                     f"net profit margin moves from {npm[0] * 100:.1f}% to "
                     f"{npm[-1] * 100:.1f}% of sales.")
    eb = [v for v in series.get("EBITDA", []) if isinstance(v, (int, float))]
    if rev and eb and rev[0]:
        lines.append(f"In Year 1, {100 - eb[0] / rev[0] * 100:.0f}% of every rupee of revenue "
                     f"is absorbed by operating costs and {eb[0] / rev[0] * 100:.0f}% remains "
                     f"as EBITDA.")
    segs = [s for s in (segments or []) if s.get("name")]
    if segs:
        top = max(segs, key=lambda s: s.get("share") or 0)
        lines.append(f"Revenue is spread across {len(segs)} target-market segments, led by "
                     f"{top['name']} at {_pct(top.get('share'))} of sales.")
    if not lines:
        return
    _table_caption(doc, "What the charts show")
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(9.5)


def _assumptions_table(doc, rows, numbering):
    """The inputs the whole model rests on — what a reviewer checks first."""
    if not rows:
        return
    _heading(doc, "Basis of Preparation — Key Assumptions", numbering,
             sub="Every projection in this report is computed from these inputs",
             new_page=True)
    table = doc.add_table(rows=1, cols=2)
    _borders(table)
    hdr = table.rows[0].cells
    _cell_text(hdr[0], "Assumption", bold=True, white=True, size=9.5)
    _shade(hdr[0], INDIGO_HEX)
    _cell_text(hdr[1], "Value", bold=True, white=True, size=9.5,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade(hdr[1], INDIGO_HEX)
    band = 0
    for label, value in rows:
        cells = table.add_row().cells
        if value is None:
            _cell_text(cells[0], label, bold=True, size=9, color=INDIGO_DEEP)
            for c in cells:
                _shade(c, "E4E4FA")
            continue
        _cell_text(cells[0], label, size=9)
        # An assumption's "value" is usually a figure, but some are words ("Straight
        # line", "Proprietorship") and those read badly pushed to the right.
        _cell_text(cells[1], value, size=9,
                   align=(FIGURE_ALIGN if _NUMERIC.match(str(value).strip())
                          else WD_ALIGN_PARAGRAPH.LEFT))
        for c in cells:
            _shade(c, LAVENDER if band % 2 == 0 else LAVENDER_SOFT)
        band += 1
    doc.add_paragraph()


def _statement_tables(doc, tables, numbering):
    """Every statutory statement the workbook computes, each on its own page."""
    for t in tables or []:
        _heading(doc, t["title"], numbering, sub=t.get("subtitle"), new_page=True)
        _table_caption(doc, f"{t['title']} (₹)")     # numbered -> List of Tables
        body = []
        for label, vals, is_heading in t["rows"]:
            if is_heading:
                body.append(([label, "", "", "", "", ""], True))
            else:
                body.append(([label] + [_inr(v) if abs(v or 0) >= 1000 else
                                        (f"{v:,.2f}" if isinstance(v, (int, float)) else "—")
                                        for v in vals], False))
        _statement_table(doc, ["Particulars", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
                         body)
        # Every statement is explained beneath it — what it is and what these numbers
        # actually do — computed from the table itself so it can never contradict it.
        try:
            from services.statement_commentary import build_commentary
            for line in build_commentary(t):
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(4)
                for r in para.runs:
                    r.font.size = Pt(9.5)
        except Exception:
            logger.warning("commentary failed for %s", t.get("key"), exc_info=True)


def _statement_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    _borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell_text(hdr[i], h, bold=True, white=True, size=9,
                   align=None if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        _shade(hdr[i], INDIGO_HEX)
    band = 0
    for cells_text, is_heading in rows:
        cells = table.add_row().cells
        if is_heading:
            _cell_text(cells[0], cells_text[0], bold=True, size=9, color=INDIGO_DEEP)
            for c in cells:
                _shade(c, "E4E4FA")
            continue
        for i, val in enumerate(cells_text):
            _cell_text(cells[i], val, size=9,
                       align=None if i == 0 else FIGURE_ALIGN)
        for c in cells:
            _shade(c, LAVENDER if band % 2 == 0 else LAVENDER_SOFT)
        band += 1
    doc.add_paragraph()


def _agent_section(doc, title, text, numbering, *, sub=None):
    """Render an agent's analysis (market research / feasibility) as a real section.
    Markdown headings from the agent become bold sub-headings rather than raw '###'."""
    if not text or not str(text).strip():
        return
    _heading(doc, title, numbering, sub=sub, new_page=True)

    def clean(s):
        # agents write markdown; the reader should never see its punctuation
        return s.replace("**", "").replace("`", "").strip()

    def flush_table(buf):
        """A run of pipe rows becomes a real table, not raw '| a | b |' text."""
        rows = []
        for r in buf:
            cells = [clean(c) for c in r.strip().strip("|").split("|")]
            if cells and not all(set(c) <= set("-: ") for c in cells):
                rows.append(cells)
        if len(rows) < 2:
            for r in rows:
                doc.add_paragraph(" — ".join(r))
            return
        width = max(len(r) for r in rows)
        body = [r + [""] * (width - len(r)) for r in rows[1:]]
        numeric = _numeric_columns(range(width), body)
        table = doc.add_table(rows=1, cols=width)
        _borders(table)
        for i in range(width):
            _cell_text(table.rows[0].cells[i], rows[0][i] if i < len(rows[0]) else "",
                       bold=True, white=True, size=9,
                       align=(FIGURE_ALIGN if numeric[i]
                              else None if i == 0 else WD_ALIGN_PARAGRAPH.LEFT))
            _shade(table.rows[0].cells[i], INDIGO_HEX)
        for bi, r in enumerate(rows[1:]):
            cells = table.add_row().cells
            for i in range(width):
                _cell_text(cells[i], r[i] if i < len(r) else "", size=9,
                           align=(FIGURE_ALIGN if numeric[i]
                                  else None if i == 0 else WD_ALIGN_PARAGRAPH.LEFT))
            for c in cells:
                _shade(c, LAVENDER if bi % 2 == 0 else LAVENDER_SOFT)
        doc.add_paragraph()

    # As in _render_markdownish: a sub-heading is bound to the next two paragraphs so it
    # cannot sit at the foot of a page with a single orphan line under it, and the run is
    # cut before the following heading so the chain stays short enough to be honoured.
    pending, skipping, hold, held = [], False, 0, None
    for raw in str(text).split("\n"):
        line = raw.strip()
        if line.startswith("|") and line.count("|") >= 2:
            if not skipping:
                pending.append(line)
            continue
        if pending:
            _cut_hold(held)
            flush_table(pending)
            pending = []
            hold, held = 0, None        # a table already anchors the heading above it
        if not line or set(line) <= set("-*_= "):
            continue                                   # blank line / horizontal rule
        if line.startswith("#"):
            head = clean(line.lstrip("#"))
            if not head:
                continue
            # A report has ONE conclusion, at the end. The agents each like to sign off
            # with their own "Conclusion & Recommendations", which put three of them in
            # the contents. Drop those sub-sections and everything under them; the
            # report-level Conclusion still comes from the model.
            bare = re.sub(r"^[\d.\s]+", "", head).strip().lower()
            skipping = bare.startswith(("conclusion", "recommendation",
                                        "summary & recommendation", "final thoughts"))
            if skipping:
                continue
            hashes = len(line) - len(line.lstrip("#"))
            _cut_hold(held)
            _heading(doc, head, level=3 if hashes >= 3 else 2)   # real sub-heading -> Contents
            hold, held = 2, None
        elif skipping:
            continue
        else:
            if line.startswith(("- ", "* ", "• ")):
                p = doc.add_paragraph(clean(line[2:]), style="List Bullet")
            else:
                p = doc.add_paragraph(clean(line))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
            held = None
            if hold:
                p.paragraph_format.keep_with_next = True
                held = p
                hold -= 1
    _cut_hold(held)
    if pending:
        flush_table(pending)


def _about_company(doc, project, numbering, *, new_page=False):
    """The borrower in plain facts — name, what it does, where, who runs it.

    Every value here comes from the client's own project record, so nothing in this
    section is a model's guess.
    """
    desc = str(project.get("project_description") or "").strip()
    answers = project.get("purpose_answers") or {}
    bank = project.get("bank_name") or answers.get("bank_name")
    facts = [
        ("Name of the unit", project.get("title")),
        ("Constitution / promoter", project.get("promoter_name")),
        ("Promoter's experience", project.get("promoter_experience")),
        ("Line of activity", project.get("sub_industry") or project.get("industry")),
        ("Industry", project.get("industry")),
        ("Location", project.get("location") or project.get("country")),
        ("Target market", project.get("target_market")),
        ("Customers served", project.get("target_customers")),
        ("Banker", bank),
    ]
    facts = [(k, str(v).strip()) for k, v in facts if v and str(v).strip()]
    if not facts and not desc:
        return

    _heading(doc, "About the Company", numbering,
             sub="The borrower, its activity, its market and what is being sought",
             new_page=new_page)
    if desc:
        p = doc.add_paragraph(desc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)

    # The proposal stated in one sentence, composed from the stored figures. This section
    # ran to under half a page of facts; a reader coming to it straight after the summary
    # should be able to see what is being asked for without turning to the annexures. Every
    # number here is the project record's own — none of it is written by a model.
    intro = _proposal_sentence(project, bank)
    if intro:
        p = doc.add_paragraph(intro)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)

    if facts:
        _table_caption(doc, "Business Particulars")
        _data_table(doc, ["Particulars", "Details"], [[k, v] for k, v in facts])

    money = _means_of_finance(project)
    if money:
        _table_caption(doc, "Cost of the Project and Means of Finance (₹)")
        _data_table(doc, ["Particulars", "Amount", "Share of project cost"], money)


def _num(v):
    """A project field as a number — these arrive as Decimal, str or None."""
    try:
        return float(str(v).replace(",", "").strip())
    except (TypeError, ValueError, AttributeError):
        return None


def _proposal_sentence(project, bank=None):
    """What is being set up, where, at what cost and funded how — from stored fields only."""
    what = str(project.get("sub_industry") or project.get("industry") or "").strip().lower()
    where = str(project.get("location") or project.get("country") or "").strip()
    cost, own, loan = (_num(project.get(k)) for k in
                       ("project_cost", "own_contribution", "loan_amount"))
    if not (cost or loan):
        return ""
    bits = [f"{project.get('title') or 'The unit'} is a {what} enterprise" if what
            else f"{project.get('title') or 'The unit'} is the borrower"]
    if where:
        bits[0] += f" at {where}"
    if project.get("promoter_name"):
        bits[0] += f", promoted by {project['promoter_name']}"
    s = bits[0] + "."
    if cost:
        s += f" The project is estimated to cost {_inr(cost)}"
        if own:
            share = f" ({own / cost * 100:.0f}% of the cost)" if cost else ""
            s += f", of which {_inr(own)}{share} is the promoter's own contribution"
        if loan:
            s += f" and {_inr(loan)} is the facility sought"
            if bank:
                s += f" from {bank}"
        s += "."
    elif loan:
        s += f" The facility sought is {_inr(loan)}"
        s += f" from {bank}." if bank else "."
    if project.get("target_market"):
        s += (f" The unit sells to {str(project['target_market']).strip().rstrip('.')}"
              f"{'' if not project.get('target_customers') else ', principally ' + str(project['target_customers']).strip().rstrip('.')}.")
    return s


def _means_of_finance(project):
    """Cost of project against how it is funded — the table a credit officer opens first."""
    cost, own, loan = (_num(project.get(k)) for k in
                       ("project_cost", "own_contribution", "loan_amount"))
    if not cost:
        return []
    rows = [["Total cost of the project", _inr(cost), "100%"]]
    for label, v in (("Promoter's contribution", own), ("Facility sought (loan)", loan)):
        if v:
            rows.append([label, _inr(v), f"{v / cost * 100:.1f}%"])
    if own and loan:
        rows.append(["Debt-equity ratio", f"{loan / own:.2f} : 1", "—"])
    return rows if len(rows) > 1 else []


def _financial_tables(doc, summary, numbering):
    series = summary.get("series") or {}
    rows = [(n, v) for n, v in series.items() if any(x is not None for x in v)]
    if rows:
        _heading(doc, "Financial Summary", numbering,
                 sub="Five-year projection — figures taken directly from the Excel model")
        _table_caption(doc, "Projected Financial Performance (₹)")
        _data_table(doc, ["Particulars", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
                    [[n] + [_inr(v) for v in vals] for n, vals in rows])

    ratios = summary.get("ratios") or {}
    rrows = [(n, v) for n, v in ratios.items() if any(x is not None for x in v)]
    if rrows:
        _heading(doc, "Key Financial Ratios", numbering)
        _table_caption(doc, "Viability and Coverage Ratios")
        body = []
        for name, vals in rrows:
            fmt = _pct if "Margin" in name else _ratio
            body.append([name] + [fmt(v) for v in vals])
        _data_table(doc, ["Ratio", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"], body)


def _segments(doc, segments, numbering):
    rows = [s for s in (segments or []) if s.get("name")]
    if not rows:
        return
    _heading(doc, "Revenue by Target Market Segment", numbering,
             sub="Where the revenue comes from — the client's own market, split by share")
    _table_caption(doc, "Revenue Contribution by Segment")
    _data_table(doc, ["Target market segment", "Share", "Year 1", "Year 5"],
                [[s["name"], _pct(s.get("share")), _inr(s.get("y1")), _inr(s.get("y5"))]
                 for s in rows])


def _swot(doc, swot, numbering):
    quads = [("Strengths", swot.get("strengths")), ("Weaknesses", swot.get("weaknesses")),
             ("Opportunities", swot.get("opportunities")), ("Threats", swot.get("threats"))]
    quads = [(t, c) for t, c in quads if c]
    if not quads:
        return
    _heading(doc, "SWOT Analysis", numbering)
    table = doc.add_table(rows=2, cols=2)
    _borders(table)
    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for (title, content), (r, c) in zip(quads, positions):
        cell = table.cell(r, c)
        cell.text = ""
        _shade(cell, LAVENDER if (r + c) % 2 == 0 else LAVENDER_SOFT)
        h = cell.paragraphs[0].add_run(title.upper())
        h.bold = True
        h.font.size = Pt(10)
        h.font.color.rgb = INDIGO_DEEP
        for line in str(content).split("\n"):
            line = line.strip().lstrip("•").strip()
            if line:
                bp = cell.add_paragraph(line, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                for run in bp.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def _conclusion(doc, text, numbering):
    _heading(doc, "Conclusion", numbering)
    table = doc.add_table(rows=1, cols=1)
    _borders(table)
    cell = table.cell(0, 0)
    _shade(cell, LAVENDER)
    cell.text = ""
    for para in str(text).split("\n"):
        para = para.strip()
        if not para:
            continue
        p = cell.paragraphs[0] if not cell.paragraphs[0].runs else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(para)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
    doc.add_paragraph()


_REF_KIND_ORDER = {"report": 0, "statistics": 0, "regulation": 0, "article": 1,
                   "website": 1, "book": 2}


def _references(doc, refs, numbering):
    """Closing References section — where the non-financial content came from.

    Each entry is a dict: kind, title, publisher, year, url (optional), note (optional).
    Reachability of every URL is checked before this point; an entry whose link did not
    respond is printed WITHOUT a link rather than with one that fails in front of a banker.
    """
    refs = [r for r in (refs or []) if isinstance(r, dict) and r.get("title")]
    if not refs:
        return
    _heading(doc, "References", numbering,
             sub="Sources for the market, industry and regulatory content of this report")

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "The financial projections in this report are computed from the promoter's own "
        "inputs and are not drawn from any external source. The market, industry and "
        "regulatory content draws on the publications below; online sources were current "
        "at the time of preparation.")
    ir.italic = True
    ir.font.size = Pt(9)
    ir.font.color.rgb = GREY
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    refs = sorted(refs, key=lambda r: (_REF_KIND_ORDER.get(
        str(r.get("kind") or "").lower(), 1), str(r.get("year") or ""), r["title"]))
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)   # hanging indent, as a list reads
        n = p.add_run(f"{i}.  ")
        n.bold = True
        n.font.size = Pt(9.5)
        n.font.color.rgb = INDIGO_DEEP

        t = p.add_run(str(r["title"]))
        t.bold = True
        t.font.size = Pt(9.5)
        t.font.color.rgb = INK

        bits = [str(r[k]) for k in ("author", "publisher", "year") if r.get(k)]
        if bits:
            m = p.add_run(". " + ", ".join(bits) + ".")
            m.font.size = Pt(9.5)
            m.font.color.rgb = GREY
        if r.get("note"):
            nr = p.add_run(" " + str(r["note"]))
            nr.italic = True
            nr.font.size = Pt(9)
            nr.font.color.rgb = GREY
        if r.get("url"):
            p.add_run(" ").font.size = Pt(9.5)
            lr = _external_link(p, str(r["url"]), str(r["url"]))
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = INDIGO
            lr.underline = True
    doc.add_paragraph()


def report_sections(model: dict, purpose_key: str, project: dict) -> list:
    """The report's main section titles, in order.

    Composed rather than listed: which sections a report actually raises depends on the
    purpose, on what the model wrote and on what the workbook produced, so any second list
    of titles would drift out of step with the document. `_compose` does no I/O and makes
    no model call, so running it purely to read the headings back is cheap and cannot
    disagree with the document the client will download.
    """
    doc = _compose(model, purpose_key, project, toc=None, skip_toc=True)
    return list(getattr(doc, "_section_titles", []))


# ── main ────────────────────────────────────────────────────────────────────────
def build_word(model: dict, purpose_key: str, project: dict) -> bytes:
    """The report, with front matter that carries real page numbers.

    Three passes, because the front matter changes the pagination it is describing:
      1. build with no front matter at all — this only tells us WHICH headings and
         captions exist;
      2. build with the three lists laid out but the numbers blank, render it, and read
         the page each heading and caption landed on;
      3. build again with those numbers typed in. Pass 3 has the same number of lines as
         pass 2, so the pagination it measured still holds.
    If the render is unavailable the lists are still written, just without numbers —
    better than the numbers all reading 1, which is what the live Word field produced.
    """
    first = _compose(model, purpose_key, project, toc=None, skip_toc=True)
    headings = _heading_list(first)
    captions = _caption_list(first)

    # SEQ numbers each kind in document order, which is the order they were collected in,
    # so the lists can say "Table 3" without the field ever being rendered.
    seq, labels = {}, []
    for kind, title, _ in captions:
        seq[kind] = seq.get(kind, 0) + 1
        labels.append(f"{kind} {seq[kind]}: {title}")

    def _cap_rows(pages):
        """The two caption lists, split by kind but keeping their document-order pages."""
        out = {"Figure": [], "Table": []}
        for i, (kind, _, anchor) in enumerate(captions):
            out[kind].append((1, labels[i], pages[i], anchor))
        return out["Figure"], out["Table"]

    def _toc_rows(pages):
        return [(lvl, text, pages[i], anchor)
                for i, (lvl, text, anchor) in enumerate(headings)]

    none_h, none_c = [None] * len(headings), [None] * len(captions)
    fig_rows, tab_rows = _cap_rows(none_c)
    draft = _docx_bytes(_compose(model, purpose_key, project, toc=_toc_rows(none_h),
                                 figures=fig_rows, tables=tab_rows))
    head_pages, cap_pages = _measure_pages(draft, headings, captions)

    fig_rows, tab_rows = _cap_rows(cap_pages)
    return _docx_bytes(_compose(model, purpose_key, project, toc=_toc_rows(head_pages),
                                figures=fig_rows, tables=tab_rows))


def _docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _compose(model: dict, purpose_key: str, project: dict, toc=None, skip_toc=False,
             figures=None, tables=None):
    config = get_config(purpose_key)
    doc = Document()
    _set_base_font(doc)
    _page_footer(doc, title=str(project.get("title") or "Project Report"),
                 logo=(project.get("branding") or {}).get("logo"))

    summary = model.get("financial_summary") or {}
    cards = summary.get("cards") or {}
    num = _Numbering()

    _set_inserts(doc, project.get("section_inserts"))
    _cover(doc, config, project)
    if not skip_toc:
        _toc_pages(doc, toc, figures, tables)

    narrative = model.get("narrative") or {}

    # 1. The executive summary opens the report — a reader should meet the case for the
    #    project before any table, not after the financial annexures.
    # The summary and the headline numbers are ONE section, not two. Splitting them put
    # "Executive Financial Snapshot" in the contents as a separate item, when the figures
    # are simply the evidence for the argument the paragraph above just made.
    exec_summary = narrative.get("Executive Summary")
    if exec_summary:
        _heading(doc, "Executive Summary", num, new_page=True,
                 sub="The case for the project and the headline numbers behind it")
        _render_markdownish(doc, exec_summary)
        _heading(doc, "Headline Numbers and Five-Year Trajectory", level=2)
    else:
        _heading(doc, "Executive Financial Snapshot", num,
                 sub="The headline numbers and their five-year trajectory")
    if cards:
        _kpi_cards(doc, cards)
    if not summary:
        kpis = model.get("kpis") or []
        if kpis:
            _data_table(doc, ["Indicator", "Value"],
                        [[k.get("label", ""), k.get("value", "")] for k in kpis[:16]])

    # Who the borrower is, in plain terms, before the analysis starts. The executive
    # summary argues the case; a reader still needs the plain facts of the business —
    # what it does, where, run by whom — and those are the client's own fields, not
    # anything the model invented.
    _about_company(doc, project, num)

    # How the business actually makes its money — the section a banker reads to understand
    # what they are lending against, before any projection makes sense.
    business_model = narrative.get("Business Model")
    if business_model:
        _heading(doc, "Business Model", num,
                 sub="How the business earns, what it costs to serve, and why it holds")
        _render_markdownish(doc, business_model)

    # 3. Market research — the agents already produce this for every report; it used to
    #    feed only the prompt and never reached the reader.
    _agent_section(doc, "Market Research & Industry Analysis",
                   model.get("market_research"), num,
                   sub="Demand, competition and the operating environment for this business")
    _agent_section(doc, "Feasibility Assessment", model.get("feasibility_analysis"), num,
                   sub="Technical, commercial and financial viability of the proposal")

    if summary:
        _financial_tables(doc, summary, num)
    _segments(doc, model.get("market_segments"), num)

    # Sections already placed at the front of the report; rendering them again here would
    # duplicate them and put two of each in the contents.
    placed = {"Executive Summary"} if exec_summary else set()
    if business_model:
        placed.add("Business Model")

    # Remaining narrative sections (the executive summary is already above). A section with
    # nothing written for it normally gets a pointer to the workbook — but on a
    # workbook-only run NOTHING is written, and six headings each followed by "refer to the
    # Excel model" (with "Executive Summary" among them, stranded in the middle of the
    # report) reads as a broken document rather than a deliberate one. So when the whole
    # narrative is absent those sections are simply not raised.
    # "Did the expensive narrative call actually run?" — measured over the standard sections
    # MINUS the ones written separately by their own cheap agents and already placed at the
    # front. Testing the whole narrative was wrong twice over: the back-fill writes
    # "Business Model" and now the "Executive Summary" into it, and each time the report
    # stopped looking empty and all five placeholder headings came back.
    any_narrative = any(str(narrative.get(s["title"]) or "").strip()
                        for s in config["word_sections"] if s["title"] not in placed)
    for section in config["word_sections"]:
        title = section["title"]
        content = narrative.get(title)
        if title in placed or (not content and not any_narrative):
            continue
        _heading(doc, title, num)
        if content:
            _render_markdownish(doc, content)
        else:
            p = doc.add_paragraph("Refer to the accompanying Excel financial model for detailed figures.")
            p.runs[0].italic = True

    # Anything the model wrote because the user asked for it by name. These have no entry
    # in word_sections, so without this loop the section was produced and then thrown away.
    standard = {s["title"] for s in config["word_sections"]} | placed
    for title, content in narrative.items():
        if title in standard or not str(content).strip():
            continue
        _heading(doc, str(title), num)
        _render_markdownish(doc, str(content))

    # Every statement the workbook computes — this is what makes it a full appraisal
    # document, and every figure is read from that same workbook.
    _assumptions_table(doc, model.get("key_assumptions"), num)
    _statement_tables(doc, model.get("statement_tables"), num)

    _swot(doc, model.get("swot") or {}, num)
    conclusion = model.get("conclusion")
    if conclusion:
        _conclusion(doc, conclusion, num)

    # Charts last, together, each on its own block.
    if summary:
        _charts_annexure(doc, summary, num, model.get("market_segments"))

    # References close the report — the last page, as asked.
    _references(doc, model.get("references"), num)
    _flush_inserts(doc)          # whatever belongs to the last section

    note = doc.add_paragraph()
    nr = note.add_run("Detailed month-by-month projections, the dashboard, ratios and statutory "
                      "formats are provided in the accompanying Excel workbook. All figures in this "
                      "report are drawn from that model.")
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = GREY

    # The closing assurance the client asked for — say plainly how the report was made,
    # because a banker's first question about any modern document is whether a model
    # simply wrote it.
    stmt = doc.add_paragraph()
    stmt.paragraph_format.space_before = Pt(10)
    _bottom_rule(stmt)
    sr = stmt.add_run(
        "This report is not generated wholly by artificial intelligence. Every figure is "
        "computed by a structured financial model from the promoter's own inputs and is "
        "reproduced here directly from that workbook; the projections follow standard "
        "CMA formats and have been checked for internal consistency. The document is "
        "complete and ready for submission to the bank.")
    sr.bold = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = INK
    stmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Only the footer page number is a field now — the contents is typed out — so Word is
    # asked to refresh fields on open without the reader being prompted about the contents.
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)

    return doc

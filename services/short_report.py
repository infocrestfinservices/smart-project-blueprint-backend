"""The SHORT report — for someone deciding whether the project is worth doing at all.

The long report is a bank submission: every CMA form, every schedule, 50-odd pages. Someone
who has not got that far wants four things — what it will cost, how it would be funded, what
the borrowing costs, and what the business should earn — and wants them in a couple of
minutes, not a couple of hundred rows.

Nothing here is computed a second time. Every figure is read from the SAME stored model the
long report and the workbook are built from, so a client who later asks for the full report
finds identical numbers. Where a figure genuinely is not on file the row is dropped rather
than estimated: a short report that quietly invents a cost is worse than one that is short.

Two deliverables:
  build_short_excel()  ->  a 2-sheet workbook: Budget & Loan, Projections
  build_short_word()   ->  a 3-page document:  Overview, Budget & Funding, What It Earns
"""
import io
import logging
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

INDIGO_HEX = "4F46E5"
LAVENDER_HEX = "EEF0FE"
SOFT_HEX = "F7F8FF"
INK = RGBColor(0x1F, 0x23, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x8A)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
INDIGO_DEEP = RGBColor(0x37, 0x30, 0xA3)

# What counts as a cost of the project rather than a running cost. Matched STRUCTURALLY —
# every industry template names its capital rows "<asset> — cost" — rather than by a list of
# asset words. A vocabulary list looked fine on the manufacturing templates and then found
# nothing at all on the software one, whose single capital row is "App Development Cost /
# Capitalised Software — cost".
_CAPEX_SUFFIX = re.compile(r"[—–-]\s*cost\s*$", re.I)
_CAPEX_WORDS = re.compile(
    r"^\s*(land|building|plant|machinery|furniture|equipment|vehicle|computer|"
    r"pre-?operative|preliminar|contingenc|security deposit|margin money|"
    r"working capital margin|civil work)\b", re.I)
_NOT_CAPEX = re.compile(r"depreciat|escalat|per unit|/ ?month|holding|cycle|"
                        r"gross margin|utilisation|tenure|moratorium", re.I)

# Cost lines that move with volume. Everything else in the cost block is treated as fixed
# for the break-even, and the report says so — a reader can disagree with the split only if
# they can see it.
_VARIABLE = re.compile(r"raw material|power|fuel|other variable|selling|distribution|"
                       r"consumable|packing|freight", re.I)


# ── reading the stored model ────────────────────────────────────────────────────
def _num(v):
    """A number out of whatever the model stored — 4,000,000 / '12.00%' / '₹1.5 L' / None."""
    if isinstance(v, (int, float)):
        return float(v)
    if not isinstance(v, str):
        return None
    t = v.strip().replace(",", "").replace("₹", "")
    pct = t.endswith("%")
    t = t.rstrip("%").strip()
    try:
        n = float(t)
    except ValueError:
        return None
    return n / 100 if pct else n


def _assumption(rows, *patterns):
    """The value of the first assumption whose label matches, or None."""
    for pat in patterns:
        rx = re.compile(pat, re.I)
        for row in rows or []:
            if len(row) >= 2 and row[0] and rx.search(str(row[0])):
                return row[1]
    return None


def _series(model, name_pattern):
    """A five-year row out of the consolidated summary, by label.

    Rows whose values are ALL empty are skipped: the summary's section headers
    ("PRODUCTION & SALES") carry a row of Nones, and matching one of those returned five
    blanks — which is why the volume line printed its label and no figures at all.
    """
    rx = re.compile(name_pattern, re.I)
    for table in (model.get("statement_tables") or []):
        for row in (table.get("rows") or []):
            if not (row and row[0] and rx.search(str(row[0])) and isinstance(row[1], list)):
                continue
            if any(isinstance(v, (int, float)) for v in row[1]):
                return list(row[1])
    return []


def _series_label(model, name_pattern):
    """The label of the row `_series` would return — the templates name it differently."""
    rx = re.compile(name_pattern, re.I)
    for table in (model.get("statement_tables") or []):
        for row in (table.get("rows") or []):
            if not (row and row[0] and rx.search(str(row[0])) and isinstance(row[1], list)):
                continue
            if any(isinstance(v, (int, float)) for v in row[1]):
                return str(row[0]).strip()
    return ""


def _table(model, title_pattern):
    rx = re.compile(title_pattern, re.I)
    for table in (model.get("statement_tables") or []):
        if rx.search(str(table.get("title") or "")):
            return table
    return None


def _volume_row(model):
    """(label, five-year values) for how much the business actually DOES each year.

    Found by POSITION, not by name: it is the first row carrying figures under the
    summary's "PRODUCTION & SALES" heading. Every industry names it differently —
    "Production (units)", "Purchase (units)", "Rooms sold / covers served",
    "Students / batches served", "Service delivery (units)" — and matching on vocabulary
    picked up "Purchases (raw material)", which is money, not units.
    """
    table = _table(model, r"Consolidated Five-Year Summary")
    seen = False
    for row in ((table or {}).get("rows") or []):
        label = str(row[0] or "")
        if re.match(r"^\s*PRODUCTION\s*&\s*SALES", label, re.I):
            seen = True
            continue
        if not seen:
            continue
        if re.match(r"^\s*(COST STRUCTURE|PROFITABILITY)", label, re.I):
            break
        if isinstance(row[1], list) and any(isinstance(v, (int, float)) for v in row[1]):
            return label.strip(), list(row[1])
    return "", []


def _cost_rows(model):
    """The operating cost lines, as (label, five-year values)."""
    out, seen = [], False
    table = _table(model, r"Consolidated Five-Year Summary")
    for row in ((table or {}).get("rows") or []):
        label = str(row[0] or "")
        if re.match(r"^\s*COST STRUCTURE", label, re.I):
            seen = True
            continue
        if re.match(r"^\s*(PROFITABILITY|PRODUCTION)", label, re.I):
            seen = False
        if seen and isinstance(row[1], list) and any(isinstance(v, (int, float)) for v in row[1]):
            out.append((label, row[1]))
    return out


def short_facts(model: dict, project: dict) -> dict:
    """Everything both deliverables need, read once out of the stored model."""
    rows = model.get("key_assumptions") or []
    summary = model.get("financial_summary") or {}
    cards = summary.get("cards") or {}

    capex = []
    for row in rows:
        label = str((row[0] if row else "") or "")
        if not label or _NOT_CAPEX.search(label):
            continue
        if not (_CAPEX_SUFFIX.search(label) or _CAPEX_WORDS.search(label)):
            continue
        value = _num(row[1] if len(row) > 1 else None)
        if value and value > 0:
            capex.append((_CAPEX_SUFFIX.sub("", label).strip(" -—–"), value))

    cost = _num(_assumption(rows, r"^TOTAL COST OF PROJECT")) or _num(project.get("project_cost"))
    # Assets coming to LESS than the stated cost is normal — the rest is the working
    # capital margin, pre-operative expenses and contingency, and the report shows it as a
    # balancing line. Assets coming to MORE is not: on one stored model they came to half
    # as much again as the project cost, which would print a table whose shares add to
    # 150%. The workbook's own reconciler scales the assets down, but this snapshot of the
    # assumptions can predate that, so the discrepancy is carried out and stated.
    capex_total = sum(v for _, v in capex) or None
    capex_ties = bool(cost and capex_total and capex_total <= cost * 1.01)
    loan = _num(_assumption(rows, r"Term Loan applied")) or _num(project.get("loan_amount"))
    own = (_num(_assumption(rows, r"Promoter.*(Capital|Equity)|Equity brought"))
           or _num(project.get("own_contribution")))

    repay = _table(model, r"Term Loan Repayment") or {}
    interest_years = []
    principal_years = []
    for row in (repay.get("rows") or []):
        label = str(row[0] or "")
        if isinstance(row[1], list):
            if re.search(r"interest", label, re.I):
                interest_years = row[1]
            elif re.search(r"principal", label, re.I):
                principal_years = row[1]

    revenue = _series(model, r"Net Sales|Revenue")
    costs = _cost_rows(model)
    # Break-even from the cost split above. Only meaningful with both sides present.
    var_y1 = sum(v[1][0] for v in costs
                 if _VARIABLE.search(v[0]) and isinstance(v[1][0], (int, float)))
    fix_y1 = sum(v[1][0] for v in costs
                 if not _VARIABLE.search(v[0]) and isinstance(v[1][0], (int, float)))
    rev_y1 = revenue[0] if revenue and isinstance(revenue[0], (int, float)) else None
    breakeven = None
    if rev_y1 and rev_y1 > var_y1:
        breakeven = fix_y1 / (rev_y1 - var_y1)          # as a share of Year-1 sales

    return {
        "title": project.get("title") or "Project",
        "activity": project.get("sub_industry") or project.get("industry") or "",
        "location": project.get("location") or project.get("country") or "",
        "promoter": project.get("promoter_name") or "",
        "description": project.get("project_description") or "",
        "years": summary.get("years") or [f"Year {i}" for i in range(1, 6)],
        "capex": capex,
        "capex_total": capex_total,
        "capex_ties": capex_ties,
        "project_cost": cost,
        "own": own,
        "loan": loan,
        "rate": _num(_assumption(rows, r"Interest rate on Term Loan")),
        "wc_rate": _num(_assumption(rows, r"Interest rate on Working Capital")),
        "tenure_months": _num(_assumption(rows, r"tenure")),
        "moratorium_months": _num(_assumption(rows, r"Moratorium")),
        "interest_years": interest_years,
        "principal_years": principal_years,
        "total_interest": sum(v for v in interest_years if isinstance(v, (int, float))) or None,
        "revenue": revenue,
        "volume": _volume_row(model),
        "costs": costs,
        "ebitda": _series(model, r"^EBITDA"),
        "pat": _series(model, r"Profit After Tax"),
        "cash": _series(model, r"Cash Accrual"),
        "dscr": (summary.get("ratios") or {}).get("DSCR") or _series(model, r"^DSCR"),
        "cards": cards,
        "breakeven": breakeven,
        "fixed_y1": fix_y1 or None,
        "variable_y1": var_y1 or None,
        "currency": project.get("currency") or "INR",
        # The agents' own verdict and market picture, already on the model — free to reuse.
        "feasibility": feasibility_digest(model.get("feasibility_analysis")),
        "market": market_digest(model.get("market_research")),
    }


# ── formatting ──────────────────────────────────────────────────────────────────
def _inr(v):
    if not isinstance(v, (int, float)):
        return "—"
    a = abs(v)
    if a >= 1e7:
        return f"₹{v / 1e7:,.2f} Cr"
    if a >= 1e5:
        return f"₹{v / 1e5:,.2f} L"
    return f"₹{v:,.0f}"


def _pct(v, places=1):
    return f"{v * 100:.{places}f}%" if isinstance(v, (int, float)) else "—"


def _ratio(v):
    return f"{v:.2f}×" if isinstance(v, (int, float)) else "—"


# ── reading the agents' own words ────────────────────────────────────────────────
# The market-research and feasibility agents run for EVERY generation, whatever the format,
# and their output is already stored on the model. Reusing it here therefore costs nothing
# and adds no time — which is the whole reason the short report can answer "should I do
# this at all" without a single extra call.
#
# What it cannot rely on is the SHAPE of that output. It is markdown written by a model, and
# across five stored projects "Recommendation" alone appeared four different ways: under its
# own heading, bolded under its own heading, ON the heading line ("## Recommendation: Needs
# Review"), and bolded with a trailing clause. One project bolded every cell of its score
# table; two wrote the scores as sub-headings instead. So everything below parses
# defensively and returns None rather than a guess — a missing block is dropped from the
# report, never filled with something invented.


def _plain(text):
    """Markdown emphasis stripped, so `**78**` and `78` parse the same."""
    return re.sub(r"[*_`]+", "", str(text or "")).strip()


def _md_table(text):
    """Rows of the first markdown table in `text`, as lists of plain cells.

    Header and separator rows are dropped: the separator is all dashes and colons, and the
    header is recognised by not carrying a number where the data rows do.
    """
    rows = []
    for line in str(text or "").split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            if rows:
                break                      # the table has ended
            continue
        cells = [_plain(c) for c in line.strip("|").split("|")]
        if not cells or all(set(c) <= set("-: ") for c in cells):
            continue                       # separator
        rows.append(cells)
    return rows


def _section(text, *patterns):
    """The body under the first heading matching any pattern, up to the next heading."""
    lines = str(text or "").split("\n")
    for i, line in enumerate(lines):
        bare = _plain(line.lstrip("#"))
        if not line.strip().startswith("#"):
            continue
        if any(re.search(p, bare, re.I) for p in patterns):
            depth = len(line) - len(line.lstrip("#"))
            out = []
            for nxt in lines[i + 1:]:
                if nxt.strip().startswith("#") and \
                        (len(nxt) - len(nxt.lstrip("#"))) <= depth:
                    break
                out.append(nxt)
            return "\n".join(out).strip(), bare
    return "", ""


def feasibility_digest(text):
    """{overall, recommendation, scores:[(category, score, assessment)]} — any part may be None."""
    text = str(text or "")
    if not text.strip():
        return {}
    out = {"overall": None, "recommendation": None, "scores": []}

    m = re.search(r"Overall\s+Feasibility\s+Score\D{0,4}(\d{1,3})\s*/\s*100", _plain(text), re.I)
    if m and 0 <= int(m.group(1)) <= 100:
        out["overall"] = int(m.group(1))

    body, heading = _section(text, r"^\s*Recommendation")
    # "## Recommendation: Needs Review" puts the answer on the heading itself.
    if ":" in heading:
        after = heading.split(":", 1)[1].strip()
        if after:
            out["recommendation"] = after
    if not out["recommendation"]:
        for line in body.split("\n"):
            line = _plain(line)
            if line:
                # keep the verdict, drop the sentence that follows it
                out["recommendation"] = re.split(r"\s+[—–-]\s+|\.\s", line)[0].strip(" .")
                break

    scores, _ = _section(text, r"Feasibility\s+Scores")
    for row in _md_table(scores):
        if len(row) < 2:
            continue
        m = re.match(r"^(\d{1,3})$", row[1].strip())
        if not m or not row[0]:
            continue
        out["scores"].append((row[0], int(m.group(1)),
                              row[2].strip() if len(row) > 2 else ""))
    if not out["scores"]:
        # some runs write them as sub-headings: "### 1. Market Potential (68/100)"
        for line in text.split("\n"):
            if not line.strip().startswith("#"):
                continue
            m = re.match(r"^[\d.\s]*(.+?)\s*[(—–-]\s*(\d{1,3})\s*/\s*100",
                         _plain(line.lstrip("#")))
            if m:
                out["scores"].append((m.group(1).strip(), int(m.group(2)), ""))
    return out


def market_digest(text):
    """{size:[(label, line)], competitors:[(name, note)]} from the stored market research."""
    text = str(text or "")
    if not text.strip():
        return {}
    out = {"size": [], "competitors": []}

    # TAM / SAM / SOM and the growth rate, wherever and however they are written.
    wanted = (("TAM", r"\bTAM\b|Total Addressable"),
              ("SAM", r"\bSAM\b|Serviceable Addressable"),
              ("SOM", r"\bSOM\b|Serviceable Obtainable"),
              ("Growth", r"CAGR|Growth Rate"))
    lines = [_plain(l) for l in text.split("\n")]
    for label, pat in wanted:
        for i, line in enumerate(lines):
            if not re.search(pat, line, re.I) or line.startswith("|"):
                continue
            # the figure may be on this line, or on the first real line beneath it
            body = re.split(r":\s*", line, 1)[1] if ":" in line else ""
            if len(body) < 15:
                body = next((x for x in lines[i + 1:i + 4]
                             if len(x) > 20 and not x.startswith(("|", "#"))), body)
            body = body.strip(" -*")
            if len(body) > 15:
                out["size"].append((label, body[:200]))
                break

    comp, _ = _section(text, r"Competitor")
    rows = _md_table(comp)
    if rows:
        head = [c.lower() for c in rows[0]]
        # the header row has no numbers; treat it as one only if it names a column
        start = 1 if any(h.startswith(("competitor", "player", "brand", "name"))
                         for h in head) else 0
        note_col = next((i for i, h in enumerate(head) if "strength" in h), None)
        for row in rows[start:]:
            name = row[0].strip()
            if not name or len(name) > 60:
                continue
            note = (row[note_col].strip() if note_col is not None and note_col < len(row)
                    else (row[1].strip() if len(row) > 1 else ""))
            out["competitors"].append((name, note[:120]))
    return out


def _explain_funding(f):
    """What the split between own money and borrowing means, in the client's own figures.

    Written from the numbers rather than by a model: the short report exists to be quick,
    and an explanation that can be derived cannot be hallucinated either.
    """
    cost, own, loan = f["project_cost"], f["own"], f["loan"]
    if not (cost and own and loan):
        return ""
    share = own / cost
    verdict = ("comfortably above the quarter of the cost a lender normally expects the "
               "promoter to fund" if share >= 0.25 else
               "below the quarter of the cost a lender normally expects the promoter to "
               "fund, which is usually the first thing questioned")
    return (f"A project of this kind is paid for from two pockets. {_inr(own)} — "
            f"{_pct(share, 0)} of the {_inr(cost)} it needs — comes from the promoter, and "
            f"{_inr(loan)} is borrowed. That promoter's share is {verdict}. The borrowing "
            f"also fixes the obligation: it has to be repaid whether or not the sales come, "
            f"which is why the coverage on the next page matters more than the profit.")


def _explain_loan(f):
    """What the interest actually costs, and what a moratorium is."""
    if not (f["loan"] and f["total_interest"]):
        return ""
    over = f["total_interest"] / f["loan"]
    bits = [f"Borrowing {_inr(f['loan'])} does not cost {_inr(f['loan'])}. Over the life of "
            f"the loan the interest adds {_inr(f['total_interest'])} — about "
            f"{_pct(over, 0)} of the amount borrowed — so the real cost of the money is "
            f"{_inr(f['loan'] + f['total_interest'])}."]
    if f["moratorium_months"]:
        bits.append(f"The first {f['moratorium_months']:.0f} months are a moratorium: no "
                    f"principal is repaid while the business is finding its feet, though "
                    f"interest still runs. Repayment of the principal begins after that, "
                    f"which is why the first year's outgo is smaller than the years after it.")
    return " ".join(bits)


def _explain_cost(f):
    """What the cost of a project is made of."""
    if not f["capex"]:
        return ""
    biggest = max(f["capex"], key=lambda x: x[1])
    base = f["project_cost"] if f["capex_ties"] else f["capex_total"]
    line = (f"The cost of a project is not only the machinery. It is everything needed "
            f"before the first rupee is earned — the assets below, and the money that has "
            f"to sit in stock and in customers' hands once trading starts.")
    if base:
        line += (f" Here the largest single item is {biggest[0].lower()} at "
                 f"{_inr(biggest[1])}, {_pct(biggest[1] / base, 0)} of the total.")
    return line


def _explain_earnings(f):
    """What the five-year projection is, and what actually drives it."""
    rev = [v for v in (f["revenue"] or []) if isinstance(v, (int, float))]
    if len(rev) < 2:
        return ""
    growth = (rev[-1] / rev[0]) if rev[0] else None
    line = (f"These are projections, not promises. They assume the business sells what the "
            f"plan says it will, at the prices it says it will, and pays what it says it "
            f"will for materials and people.")
    if growth:
        line += (f" On those assumptions sales grow from {_inr(rev[0])} to {_inr(rev[-1])} "
                 f"over five years — {growth:.1f} times — mostly by using more of the "
                 f"capacity that has already been paid for, which is why the profit grows "
                 f"faster than the sales do.")
    return line


def _explain_ratios(f):
    """DSCR, break-even and margin in plain words, using this project's own values."""
    out = []
    dscr = f["cards"].get("avg_dscr")
    if isinstance(dscr, (int, float)):
        out.append(f"**Loan cover (DSCR)** is the one number a lender looks at first. It "
                   f"asks: for every ₹1 the business must pay the bank in a year, how many "
                   f"rupees does it actually generate? Here it is {dscr:.2f}. Anything "
                   f"below 1.25 is usually refused; below 1.00 the business cannot pay "
                   f"from its own earnings at all.")
    if f["breakeven"]:
        out.append(f"**Break-even** is the point at which the business stops losing money. "
                   f"Its fixed costs — rent, salaries, depreciation, interest — have to be "
                   f"paid whatever it sells. Those are covered once it reaches "
                   f"{_pct(f['breakeven'])} of its year-1 sales target; every rupee after "
                   f"that is profit. The lower this is, the more room there is to be wrong "
                   f"about demand.")
    margin = f["cards"].get("net_margin_y5")
    if isinstance(margin, (int, float)):
        out.append(f"**Net margin** is what is left after everything, including tax and "
                   f"interest. At {_pct(margin)} by year five, every ₹100 of sales leaves "
                   f"₹{margin * 100:.0f} for the promoter to keep or reinvest.")
    return out


def _score_for(f, pattern):
    """One of the feasibility review's category scores, by name."""
    rx = re.compile(pattern, re.I)
    for cat, sc, _ in ((f.get("feasibility") or {}).get("scores") or []):
        if rx.search(str(cat)):
            return sc
    return None


def _verdict(f):
    """Should this business be started at all?

    NOT whether the loan is covered. The person this report is for has not settled the
    financing — they may not even know what the project will cost yet — so a verdict built
    on DSCR answers a question they have not reached. What they are deciding is whether the
    business is worth doing, and that rests on the demand, on what the review made of the
    proposition, and on whether the trading economics leave anything behind.

    Loan cover is still reported, on the earnings page, where it belongs.
    """
    fd = f.get("feasibility") or {}
    market = _score_for(f, r"market\s*potential")
    overall = fd.get("overall")
    competitive = _score_for(f, r"competitive")
    margin = f["cards"].get("net_margin_y5")
    breakeven = f.get("breakeven")

    if market is None and overall is None and not isinstance(margin, (int, float)):
        return ("There is not enough on file yet to judge whether this business is worth "
                "starting — the market review and the projections are what decide it.")

    # The demand side leads: a business with no market cannot be financed into working.
    parts = []
    if market is not None:
        parts.append(f"the market review scores demand for this business {market}/100")
    if competitive is not None:
        parts.append(f"its competitive position {competitive}/100")
    if overall is not None:
        parts.append(f"and rates the proposition {overall}/100 overall")
    lead = ", ".join(parts) if parts else ""

    economics = []
    if isinstance(margin, (int, float)):
        economics.append(f"{_pct(margin)} of sales is left as profit by year five")
    if isinstance(breakeven, (int, float)):
        economics.append(f"and the business covers its fixed costs at "
                         f"{_pct(breakeven)} of its year-one sales target")
    econ = ", ".join(economics)

    strong = ((market or 0) >= 65 and (overall or 0) >= 60
              and isinstance(margin, (int, float)) and margin > 0.08
              and (breakeven is None or breakeven < 0.85))
    weak = ((market or 100) < 50 or (overall or 100) < 50
            or (isinstance(margin, (int, float)) and margin <= 0))

    if strong:
        head = "On these figures this business looks worth starting."
        tail = ("The demand is there and the trading economics leave a real margin. What "
                "would still have to be settled is the funding and the working capital — "
                "the pages before this one set those out.")
    elif weak:
        head = "On these figures this business needs a closer look before it is started."
        tail = ("The weakness is in the proposition itself, not in the financing — more "
                "borrowing would not fix it. The price, the cost base or the market being "
                "targeted is what would have to change.")
    else:
        head = "On these figures this business is workable but not comfortable."
        tail = ("It stands up, with little room for the demand or the prices to come in "
                "below plan. Worth testing the assumptions that matter most before "
                "committing money to it.")
    middle = f" {lead[0].upper()}{lead[1:]}." if lead else ""
    if econ:
        middle += f" On the trading side, {econ}."
    return head + middle + " " + tail


# ── Excel: two sheets ───────────────────────────────────────────────────────────
_THIN = Side(style="thin", color="D8DBEC")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _head(ws, row, cells, width=None):
    for i, text in enumerate(cells, start=1):
        c = ws.cell(row=row, column=i, value=text)
        c.font = Font(bold=True, color="FFFFFF", size=10)
        c.fill = PatternFill("solid", fgColor=INDIGO_HEX)
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = _BORDER
    ws.row_dimensions[row].height = 20
    return row + 1


def _title(ws, row, text):
    c = ws.cell(row=row, column=1, value=text)
    c.font = Font(bold=True, size=13, color=INDIGO_HEX)
    return row + 2


def _band(ws, row, label, value=None, extra=None, bold=False, money=True):
    c = ws.cell(row=row, column=1, value=label)
    c.font = Font(bold=bold, size=10)
    c.alignment = Alignment(wrap_text=True, vertical="center")
    c.border = _BORDER
    v = ws.cell(row=row, column=2,
                value=(round(value, 2) if isinstance(value, (int, float)) else value))
    if isinstance(value, (int, float)) and money:
        v.number_format = '#,##0'
    v.font = Font(bold=bold, size=10)
    v.border = _BORDER
    e = ws.cell(row=row, column=3, value=extra)
    e.font = Font(bold=bold, size=10, color="6B728A")
    e.border = _BORDER
    fill = LAVENDER_HEX if bold else SOFT_HEX
    for col in (1, 2, 3):
        ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor=fill)
    return row + 1


def build_short_excel(model: dict, project: dict) -> bytes:
    """Two sheets: what it costs and what it borrows, then what it should earn."""
    f = short_facts(model, project)
    wb = Workbook()

    # ── sheet 1: Budget & Loan ──
    ws = wb.active
    ws.title = "Budget & Loan"
    ws.column_dimensions["A"].width = 44
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 30
    r = _title(ws, 1, f["title"])
    ws.cell(row=2, column=1,
            value=f"{f['activity']}{' · ' + f['location'] if f['location'] else ''}"
            ).font = Font(size=10, color="6B728A")

    r = 4
    fd = f.get("feasibility") or {}
    r = _head(ws, r, ["The business", "", ""])
    for label, value in (("Promoter", f["promoter"]), ("Line of activity", f["activity"]),
                         ("Location", f["location"])):
        if value:
            r = _band(ws, r, label, value, money=False)

    # Stated as what it COULD cost, not as a settled figure. The person reading this has
    # usually not costed the project — that is part of what they came to find out — so a
    # bare "TOTAL COST OF THE PROJECT" reads as a fact they supplied when it is an estimate
    # built from the plan.
    r += 1
    if f["project_cost"]:
        r = _band(ws, r, "This business could cost about", f["project_cost"],
                  "to set up — broken down below", bold=True)
        r += 1
    r = _head(ws, r, ["What that money goes on", "Amount", "Share"])
    base = f["project_cost"] if f["capex_ties"] else f["capex_total"]
    for label, value in f["capex"]:
        r = _band(ws, r, label, value, _pct(value / base) if base else None)
    if f["capex_ties"] and f["project_cost"] - f["capex_total"] > 1:
        rest = f["project_cost"] - f["capex_total"]
        r = _band(ws, r, "Working capital margin & other", rest,
                  _pct(rest / f["project_cost"]))
    elif f["capex"] and not f["capex_ties"]:
        r = _band(ws, r, "Total of the items above", f["capex_total"], "100%", bold=True)
    r = _band(ws, r, "TOTAL", f["project_cost"],
              "100%" if f["capex_ties"] or not f["capex"]
              else "the items above are the pre-scaling asset figures",
              bold=True)

    # The promoter's own contribution is NOT shown. It is one of the figures they have not
    # settled — the same reason the cost above is put as an estimate — and printing
    # "Promoter's own money ₹20.00 L (50%)" states as fact something nobody has decided.
    # The loan being sought is what they came with, so that stays.
    r += 1
    r = _head(ws, r, ["How it is funded", "Amount", "Share"])
    if f["loan"]:
        r = _band(ws, r, "Loan sought", f["loan"],
                  _pct(f["loan"] / f["project_cost"]) if f["project_cost"] else None)

    r += 1
    r = _head(ws, r, ["What the loan costs", "", ""])
    if f["rate"] is not None:
        r = _band(ws, r, "Interest rate on the term loan", _pct(f["rate"], 2), money=False)
    if f["tenure_months"]:
        yrs = f["tenure_months"] / 12
        r = _band(ws, r, "Repayment period", f"{f['tenure_months']:.0f} months ({yrs:.1f} years)",
                  money=False)
    if f["moratorium_months"]:
        r = _band(ws, r, "Moratorium (no principal repaid)",
                  f"{f['moratorium_months']:.0f} months", money=False)
    if f["wc_rate"] is not None:
        r = _band(ws, r, "Interest rate on working capital", _pct(f["wc_rate"], 2), money=False)
    if f["total_interest"]:
        r = _band(ws, r, "TOTAL INTEREST PAYABLE", f["total_interest"],
                  "over the years shown below", bold=True)

    if f["interest_years"] or f["principal_years"]:
        r += 1
        years = f["years"][:max(len(f["interest_years"]), len(f["principal_years"]))]
        r = _head(ws, r, ["Year by year"] + [str(y) for y in years])
        for label, vals in (("Interest", f["interest_years"]),
                            ("Principal repaid", f["principal_years"])):
            c = ws.cell(row=r, column=1, value=label)
            c.font = Font(size=10)
            c.border = _BORDER
            c.fill = PatternFill("solid", fgColor=SOFT_HEX)
            for i, v in enumerate(vals, start=2):
                cell = ws.cell(row=r, column=i,
                               value=round(v, 0) if isinstance(v, (int, float)) else v)
                cell.number_format = '#,##0'
                cell.border = _BORDER
                cell.fill = PatternFill("solid", fgColor=SOFT_HEX)
            r += 1

    # The evidence, and THEN the verdict. It sits at the foot of the sheet rather than the
    # head for the same reason it closes the Word report: it is a conclusion, and a
    # conclusion printed above its evidence asks to be taken on trust.
    if fd.get("scores"):
        r += 1
        r = _head(ws, r, ["The feasibility review", "Score /100", "Assessment"])
        for cat, sc, note_text in fd["scores"]:
            r = _band(ws, r, cat, sc, note_text, money=False)
            ws.cell(row=r - 1, column=3).alignment = Alignment(wrap_text=True,
                                                               vertical="center")
        if fd.get("overall"):
            r = _band(ws, r, "OVERALL", fd["overall"], fd.get("recommendation") or "",
                      bold=True, money=False)

    r += 1
    r = _head(ws, r, ["SHOULD THIS BUSINESS BE STARTED?", "", ""])
    if fd.get("recommendation") or fd.get("overall"):
        score = f"{fd['overall']}/100" if fd.get("overall") else ""
        r = _band(ws, r, "Feasibility review says",
                  " — ".join(x for x in (fd.get("recommendation"), score) if x),
                  bold=True, money=False)
    market = _score_for(f, r"market\s*potential")
    if market is not None:
        r = _band(ws, r, "Market potential", f"{market}/100", bold=True, money=False)
    r = _band(ws, r, "The verdict", _verdict(f), bold=True, money=False)
    ws.row_dimensions[r - 1].height = 74
    ws.cell(row=r - 1, column=2).alignment = Alignment(wrap_text=True, vertical="top")

    r += 1
    note = ws.cell(row=r, column=1,
                   value="An indicative overview. Every figure is taken from the same "
                         "financial model as the full report — nothing here is a separate "
                         "estimate. The feasibility scores are the assessment prepared for "
                         "this project; they judge the business, not the cash flow.")
    note.font = Font(italic=True, size=9, color="6B728A")

    # ── sheet 2: Projections ──
    ws2 = wb.create_sheet("Projections")
    ws2.column_dimensions["A"].width = 40
    for i in range(2, 8):
        ws2.column_dimensions[get_column_letter(i)].width = 15
    r = _title(ws2, 1, "What the business should earn")
    years = [str(y) for y in f["years"]][:5]

    def block(row, heading, lines, fmt=_inr, money=True):
        row = _head(ws2, row, [heading] + years)
        for label, vals in lines:
            if not vals:
                continue
            c = ws2.cell(row=row, column=1, value=label)
            c.font = Font(size=10)
            c.border = _BORDER
            c.fill = PatternFill("solid", fgColor=SOFT_HEX)
            for i, v in enumerate(vals[:5], start=2):
                cell = ws2.cell(row=row, column=i,
                                value=(round(v, 2) if isinstance(v, (int, float)) and money
                                       else fmt(v) if not money else v))
                if money and isinstance(v, (int, float)):
                    cell.number_format = '#,##0'
                cell.border = _BORDER
                cell.fill = PatternFill("solid", fgColor=SOFT_HEX)
            row += 1
        return row + 1

    # Revenue is NOT shown. The client did not supply a sales figure — the model
    # estimated one from the plan — so printing it invites it to be read as their own
    # projection. The volume line stays: it is what the business would have to DO.
    r = block(4, "Sales", [(f["volume"][0] or "Volume", f["volume"][1])])
    r = block(r, "What it costs to run", f["costs"])
    r = block(r, "What is left", [("EBITDA", f["ebitda"]),
                                  ("Profit after tax", f["pat"]),
                                  ("Cash generated", f["cash"])])

    r = _head(ws2, r, ["The numbers that matter", "", ""])
    cards = f["cards"]
    for label, value in (
            ("EBITDA margin (year 5)",
             _pct(cards.get("ebitda_y5") / cards["revenue_y5"])
             if cards.get("revenue_y5") and cards.get("ebitda_y5") else None),
            ("Net profit margin (year 5)", _pct(cards.get("net_margin_y5"))),
            ("Average loan cover (DSCR)", _ratio(cards.get("avg_dscr"))),
            ("Break-even — share of year-1 sales needed", _pct(f["breakeven"]))):
        if value and value != "—":
            r = _band(ws2, r, label, value, money=False)

    r += 1
    v = ws2.cell(row=r, column=1, value=_verdict(f))
    v.font = Font(bold=True, size=10)
    v.alignment = Alignment(wrap_text=True, vertical="top")
    ws2.merge_cells(start_row=r, start_column=1, end_row=r + 2, end_column=6)
    r += 4
    if f["fixed_y1"] and f["variable_y1"]:
        n = ws2.cell(row=r, column=1,
                     value="Break-even treats raw material, power, other variable costs and "
                           "selling expenses as varying with volume; everything else, "
                           "including wages, is treated as fixed.")
        n.font = Font(italic=True, size=9, color="6B728A")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Word: three pages ───────────────────────────────────────────────────────────
def _w_heading(doc, text, sub=None, first=False, anchor=None):
    p = doc.add_paragraph()
    if not first:
        p.paragraph_format.page_break_before = True
    if anchor:
        # The long report's own bookmark helper, so the contents links behave identically.
        from services.word_builder import _bookmark
        _bookmark(p, anchor, abs(hash(anchor)) % 90000 + 1000)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    # Title Case, applied here as well as in the section list, so a heading added later
    # cannot arrive in sentence case by accident. The sub-line below is left alone — it is
    # a sentence, not a heading.
    r = p.add_run(_tc(text))
    r.bold = True
    r.font.size = Pt(19)
    r.font.color.rgb = INDIGO
    r.font.name = "Calibri"
    if sub:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.keep_with_next = True
        sr = sp.add_run(sub)
        sr.italic = True
        sr.font.size = Pt(9)
        sr.font.color.rgb = GREY


def _tc(text):
    """Title Case for headings — every word's first letter capitalised.

    Only the first letter of each word is touched, so an acronym or a figure keeps its own
    shape: "THE VERDICT" stays shouting, "CMA" stays "CMA", "Year 1" stays "Year 1".
    """
    out = []
    for word in str(text or "").split(" "):
        i = next((n for n, ch in enumerate(word) if ch.isalpha()), None)
        out.append(word if i is None else word[:i] + word[i].upper() + word[i + 1:])
    return " ".join(out)


def _w_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(_tc(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, INDIGO_HEX)
    for n, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if value is None else str(value))
            run.font.size = Pt(9.5)
            # Only a total is emphasised. Testing `isupper()` also caught "EBITDA", which
            # then read as if it were a summary line.
            run.bold = str(row[0] or "").strip().upper().startswith("TOTAL")
            _shade_cell(cells[i], LAVENDER_HEX if n % 2 == 0 else SOFT_HEX)
    if widths:
        for row in table.rows:
            for i, cm in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(cm)
    doc.add_paragraph()
    return table


def _shade_cell(cell, hex_fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _w_rich(doc, text, *, size=10, after=6):
    """A paragraph where **…** is bold — the explanations lead with the term being explained."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", str(text))):
        if not part:
            continue
        run = p.add_run(part)
        run.bold = i % 2 == 1
        run.font.size = Pt(size)
        run.font.color.rgb = INK
    return p


def _w_para(doc, text, *, size=10.5, italic=False, bold=False, colour=INK, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    r.font.color.rgb = colour
    return p


# The four sections, in order — the contents lists exactly these, and the pages are
# measured against these titles. The VERDICT comes LAST, deliberately: it is a conclusion,
# and a conclusion printed before its evidence asks to be taken on trust.
SHORT_SECTIONS = [
    ("Is The Market There?",
     "The demand, the competition, and what the feasibility review made of it"),
    ("What It Costs, And How It Is Paid For",
     "Where the money goes, where it comes from, and what the borrowing costs"),
    ("What It Should Earn",
     "Five-year projection — the same model the full report is built from"),
    ("Should This Business Be Started?",
     "The verdict, on the demand and the trading economics"),
]
_ANCHORS = {title: f"_Short{i:02d}" for i, (title, _) in enumerate(SHORT_SECTIONS, 1)}


def _short_toc(doc, pages):
    """A four-line contents, linked the same way the long report's is.

    Deliberately not the long report's `_list_block`: that is pinned to a 17.6 cm grid for
    the long report's margins, and this document has its own. The look and the link
    mechanism are shared; only the width differs.
    """
    from services.word_builder import _internal_link
    h = doc.add_paragraph()
    h.paragraph_format.page_break_before = True
    h.paragraph_format.space_after = Pt(10)
    hr = h.add_run("Contents")
    hr.bold = True
    hr.font.size = Pt(19)
    hr.font.color.rgb = INDIGO
    hr.font.name = "Calibri"
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    nr = note.add_run("An indicative overview to help decide whether this project is worth "
                      "taking further — not the full CMA submission. Every figure is read "
                      "from the same financial model as the full report.")
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = GREY

    table = doc.add_table(rows=1, cols=2)
    head = table.rows[0].cells
    for cell, text, align in ((head[0], "Section", None),
                              (head[1], "Page", WD_ALIGN_PARAGRAPH.RIGHT)):
        cell.text = ""
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, INDIGO_HEX)
    for n, (title, _) in enumerate(SHORT_SECTIONS):
        cells = table.add_row().cells
        page = pages.get(title)
        for cell, text, align in ((cells[0], title, None),
                                  (cells[1], "" if page is None else str(page),
                                   WD_ALIGN_PARAGRAPH.RIGHT)):
            cell.text = ""
            p = cell.paragraphs[0]
            if align:
                p.alignment = align
            run = _internal_link(p, _ANCHORS[title], text)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            _shade_cell(cell, LAVENDER_HEX if n % 2 == 0 else SOFT_HEX)
    for row in table.rows:
        row.cells[0].width = Cm(14.4)
        row.cells[1].width = Cm(2.4)
    doc.add_paragraph()


def _w_verdict_box(doc, f):
    """The two verdicts, side by side and never reconciled.

    One is the feasibility agent's own score, the other is what the figures say about the
    loan. They do NOT always agree — on one stored project the agent said "Moderately
    Feasible" while the cash flow covered the loan three times over, and on another both
    were negative. Showing only one, or averaging them into a single cheerful line, would
    be the report deciding for the reader. This is the page that has to help them decide.
    """
    fd = f.get("feasibility") or {}
    rows = []
    if fd.get("recommendation") or fd.get("overall"):
        score = f"{fd['overall']}/100" if fd.get("overall") else ""
        verdict = " — ".join(x for x in (fd.get("recommendation"), score) if x)
        rows.append(["Feasibility assessment", verdict])
    # The market-potential score is NOT given its own row here — the verdict sentence
    # already quotes it, and a third row ran this page onto a seventh on five of seven
    # projects. It has its own row in the workbook, where space is not the constraint.
    rows.append(["The verdict", _verdict(f)])
    table = _w_table(doc, ["THE VERDICT", ""], rows, widths=[4.6, 12.4])
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
    if len(rows) > 1:
        _w_para(doc, "The two are arrived at separately — the review's own scoring of the "
                     "proposition, and the same question read off the projections — and "
                     "will not always agree. Neither judges the loan; whether the borrowing "
                     "is serviceable is on the earnings page.",
                size=8.5, italic=True, colour=GREY, after=8)


def _clip(text, n):
    """Cut at a word boundary. The agents write at length; this page has one page."""
    t = " ".join(str(text or "").split())
    if len(t) <= n:
        return t
    cut = t[:n]
    return (cut[:cut.rfind(" ")] if " " in cut else cut).rstrip(" ,;:.") + "…"


def _w_market_and_feasibility(doc, f):
    """Page two: is there a market, and what did the feasibility review score it?

    Everything here is clipped hard. The market and feasibility agents write for the FULL
    report, where they have as many pages as they need; dropped in whole they ran this page
    over onto a fifth that was a fifth full.
    """
    md, fd = f.get("market") or {}, f.get("feasibility") or {}
    if not (md.get("size") or md.get("competitors") or fd.get("scores")):
        return False
    title, sub = SHORT_SECTIONS[0]
    _w_heading(doc, title, sub=sub, anchor=_ANCHORS[title])

    if md.get("size"):
        # Three at most: with four the page ran over by a single row on one project.
        _w_table(doc, ["The market", ""],
                 [[label, _clip(line, 132)] for label, line in md["size"][:3]],
                 widths=[2.6, 14.4])
    if md.get("competitors"):
        _w_table(doc, ["Who you are up against", "Their strength"],
                 [[_clip(n, 38), _clip(note, 58)] for n, note in md["competitors"][:3]],
                 widths=[5.5, 11.5])
    if fd.get("scores"):
        rows = [[_clip(cat, 32), f"{sc}/100", _clip(note, 60)]
                for cat, sc, note in fd["scores"]]
        if fd.get("overall"):
            rows.append(["OVERALL", f"{fd['overall']}/100",
                         _clip(fd.get("recommendation") or "", 60)])
        _w_table(doc, ["Feasibility review", "Score", "Assessment"], rows,
                 widths=[4.4, 1.8, 10.8])
    return True


def build_short_word(model: dict, project: dict) -> bytes:
    """Six pages: cover, contents, then the verdict, the market, the cost and the earnings.

    Two passes, for the same reason the long report needs three: a contents page can only
    carry real page numbers if something has measured them. Pass one lays the document out
    with the numbers blank, renders it, and reads back which page each section landed on;
    pass two types them in. The contents is the same four lines either way, so the
    pagination measured in pass one still holds in pass two. No model call in either.
    """
    first = _compose_short(model, project, {})
    pages = _measure_short(first)
    return _compose_short(model, project, pages) if pages else first


def _measure_short(docx_bytes):
    """{section title: page} read off a real render. {} if that is not possible."""
    try:
        from services.recalc_service import libreoffice_available
        from services.word_builder import _docx_to_pdf
        if not libreoffice_available():
            return {}
        import fitz
        pdf = _docx_to_pdf(docx_bytes)
        if not pdf:
            return {}
        with fitz.open(stream=pdf, filetype="pdf") as d:
            texts = [d[i].get_text() for i in range(d.page_count)]
        out, start = {}, 0
        for title, _ in SHORT_SECTIONS:
            for i in range(start, len(texts)):
                # the contents page names every section; the body page STARTS with one
                head = " ".join(texts[i].strip().split("\n")[0].split())
                if head.startswith(title[:34]):
                    out[title] = i + 1
                    start = i
                    break
        return out
    except Exception:
        logger.warning("short report: could not measure the contents pages", exc_info=True)
        return {}


def _compose_short(model: dict, project: dict, pages: dict) -> bytes:
    f = short_facts(model, project)
    cards = f["cards"]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style.paragraph_format.line_spacing = 1.2
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.9)
        s.right_margin = Cm(1.9)

    # ── page 1: the cover ──
    # The LONG report's own cover, reused rather than redrawn — same band, monogram, ribbon
    # and artwork, so the two documents are visibly the same family. It reads only
    # `config["label"]`, which is what names it an overview rather than a submission.
    # The label must stay SHORT: it is printed down the 3.4 cm band as "<label> REPORT",
    # and "Indicative Overview" broke mid-word there ("INDICATIVE OVERVI / EW REPORT").
    # What this document is, and is not, is said in full on the contents page below.
    try:
        from services.word_builder import _cover
        _cover(doc, {"label": "Overview"}, project)
    except Exception:
        logger.warning("short report: cover unavailable", exc_info=True)

    # ── page 2: contents ──
    _short_toc(doc, pages)

    # ── page 3: is there a market ──
    _w_market_and_feasibility(doc, f)

    # ── page 4: budget & funding ──
    title, sub = SHORT_SECTIONS[1]
    _w_heading(doc, title, sub=sub, anchor=_ANCHORS[title])
    # This is the fullest page, and it grows when the capital figures do not tie (an extra
    # total row plus the note explaining why). The least essential paragraph is the one
    # that gives way — the note below already says what this intro would have said.
    intro = _explain_cost(f) if f["capex_ties"] else ""
    if intro:
        _w_rich(doc, intro, size=9.5)
    if f["capex"]:
        base = f["project_cost"] if f["capex_ties"] else f["capex_total"]
        crows = [[label, _inr(v), _pct(v / base) if base else "—"]
                 for label, v in f["capex"]]
        if f["capex_ties"] and f["project_cost"] - f["capex_total"] > 1:
            rest = f["project_cost"] - f["capex_total"]
            crows.append(["Working capital margin & other", _inr(rest),
                          _pct(rest / f["project_cost"])])
            crows.append(["TOTAL", _inr(f["project_cost"]), "100%"])
        else:
            crows.append(["TOTAL", _inr(f["capex_total"]), "100%"])
        _w_table(doc, ["What the money is spent on", "Amount", "Share"], crows,
                 widths=[9.5, 4.0, 3.5])
        if not f["capex_ties"] and f["project_cost"]:
            _w_para(doc,
                    f"The asset figures above are the ones entered before the model scaled "
                    f"them to the stated project cost of {_inr(f['project_cost'])}; the "
                    f"projections and the loan are built on that stated cost.",
                    size=8.5, italic=True, colour=GREY, after=10)

    # The promoter's own contribution is deliberately absent — see the note in the workbook
    # builder. Only the loan being sought is shown, because that is the figure they came
    # with rather than one the model settled for them.
    frows = []
    if f["loan"]:
        frows.append(["Loan sought", _inr(f["loan"]),
                      _pct(f["loan"] / f["project_cost"]) if f["project_cost"] else "—"])
    if frows:
        _w_table(doc, ["How it is funded", "Amount", "Share"], frows, widths=[9.5, 4.0, 3.5])
        _w_para(doc, "The balance would come from the promoter's own funds. How much of it "
                     "they put in is theirs to decide — it is not assumed here.",
                size=8.5, italic=True, colour=GREY, after=8)

    lrows = []
    if f["rate"] is not None:
        lrows.append(["Interest rate", _pct(f["rate"], 2)])
    if f["tenure_months"]:
        lrows.append(["Repayment period",
                      f"{f['tenure_months']:.0f} months ({f['tenure_months'] / 12:.1f} years)"])
    if f["moratorium_months"]:
        lrows.append(["Moratorium — no principal repaid",
                      f"{f['moratorium_months']:.0f} months"])
    if f["total_interest"]:
        lrows.append(["Total interest payable", _inr(f["total_interest"])])
    obligation = None
    if f["interest_years"] and f["principal_years"]:
        first_year = [(f["interest_years"][0] if f["interest_years"] else 0),
                      (f["principal_years"][0] if f["principal_years"] else 0)]
        if all(isinstance(v, (int, float)) for v in first_year):
            obligation = sum(first_year)
            lrows.append(["What must be paid in year 1",
                          f"{_inr(obligation)}  (about {_inr(obligation / 12)} a month)"])
    if lrows:
        _w_table(doc, ["What the loan costs", ""], lrows, widths=[9.5, 7.5])
    loan_note = _explain_loan(f)
    if loan_note:
        _w_rich(doc, loan_note, size=9.5, after=5)

    # ── page 5: what it earns ──
    title, sub = SHORT_SECTIONS[2]
    _w_heading(doc, title, sub=sub, anchor=_ANCHORS[title])
    earn = _explain_earnings(f)
    if earn:
        _w_rich(doc, earn)
    years = [str(y) for y in f["years"]][:5]
    prows = []
    for label, vals in (("EBITDA", f["ebitda"]),
                        ("Profit after tax", f["pat"]), ("Cash generated", f["cash"])):
        if vals:
            prows.append([label] + [_inr(v) for v in vals[:5]])
    if prows:
        _w_table(doc, ["Particulars"] + years, prows, widths=[4.4] + [2.6] * len(years))

    krows = []
    if cards.get("revenue_y5") and cards.get("ebitda_y5"):
        krows.append(["EBITDA margin by year 5",
                      _pct(cards["ebitda_y5"] / cards["revenue_y5"])])
    if isinstance(cards.get("net_margin_y5"), (int, float)):
        krows.append(["Net profit margin by year 5", _pct(cards["net_margin_y5"])])
    if cards.get("avg_dscr"):
        krows.append(["Average loan cover (DSCR)",
                      f"{_ratio(cards['avg_dscr'])}  (a lender usually looks for 1.25×)"])
    if f["breakeven"]:
        krows.append(["Break-even",
                      f"{_pct(f['breakeven'])} of year-1 sales covers all fixed costs"])
    if krows:
        _w_table(doc, ["The numbers that decide it", ""], krows, widths=[9.0, 8.0])
    for line in _explain_ratios(f):
        _w_rich(doc, line, size=9.5, after=5)

    if f["fixed_y1"] and f["variable_y1"]:
        _w_para(doc, "Break-even treats raw material, power, other variable costs and selling "
                     "expenses as varying with volume; everything else, including wages, is "
                     "treated as fixed.", size=8.5, italic=True, colour=GREY, after=4)

    # ── page 6: the verdict, after the evidence for it ──
    title, sub = SHORT_SECTIONS[3]
    _w_heading(doc, title, sub=sub, anchor=_ANCHORS[title])
    where = f" at {f['location']}" if f["location"] else ""
    who = f", promoted by {f['promoter']}" if f["promoter"] else ""
    _w_para(doc, f"{f['title']} is a {f['activity'].lower() or 'proposed'} venture{where}{who}.")
    if f["description"]:
        _w_para(doc, f["description"], after=10)

    _w_verdict_box(doc, f)

    rows = []
    if f["project_cost"]:
        rows.append(["Could cost about", _inr(f["project_cost"])])
    if f["loan"]:
        rows.append(["Loan sought", _inr(f["loan"])])
    cards = f["cards"]
    # No sales figure here either — see the note on the Projections sheet.
    if cards.get("pat_y5"):
        rows.append(["Profit by year 5", _inr(cards["pat_y5"])])
    if isinstance(cards.get("net_margin_y5"), (int, float)):
        rows.append(["Net margin by year 5", _pct(cards["net_margin_y5"])])
    if rows:
        _w_table(doc, ["At a glance", ""], rows, widths=[10.5, 6.5])

    _w_para(doc,
            "This is a short, indicative overview to help decide whether the project is "
            "worth taking further. Every figure is read from the same financial model as "
            "the full report, so the numbers will not change if you ask for it — but this "
            "document is not a bank submission. For that, generate the full report, which "
            "carries the CMA forms, the detailed schedules and the supporting analysis a "
            "lender requires.",
            size=9, italic=True, colour=GREY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
